import streamlit as st
from openai import OpenAI
import json, re, io, base64, subprocess, tempfile, os
from pathlib import Path

try:
    import stripe as stripe_lib
    STRIPE_AVAILABLE = True
except ImportError:
    STRIPE_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False



PRIMARY    = "#1A1A1A"
PRIMARY_MID= "#2E2E2E"
GOLD       = "#D4B36E"
GOLD_DARK  = "#A8863A"
GOLD_LIGHT = "#FAF4E6"
GOLD_MID   = "#F5EDD6"
GOLD_BG    = "#FFFBEC"
GREEN      = "#1A7A4A"
GREEN_BG   = "#D4EDDA"
NAVY       = "#1C3A5E"
NAVY_LIGHT = "#DDE8F4"
NAVY_BG    = "#EEF4FB"

# ── Helpers ───────────────────────────────────────────────────────────────────
def get_logo_b64():
    p = Path(__file__).parent / "logo.png"
    if p.exists():
        with open(p, "rb") as f:
            return base64.b64encode(f.read()).decode()
    return None

def get_secret(key, default=""):
    try:
        return st.secrets.get(key, default)
    except Exception:
        return default

def verify_stripe_session(session_id):
    if not STRIPE_AVAILABLE:
        return False, None
    try:
        stripe_lib.api_key = get_secret("STRIPE_SECRET_KEY")
        session = stripe_lib.checkout.Session.retrieve(session_id)
        if session.payment_status == "paid":
            email = session.customer_details.email if session.customer_details else None
            return True, email
    except Exception:
        pass
    return False, None

def is_authenticated():
    return (st.session_state.get("ad_access_granted", False) or
            st.session_state.get("lifetime_access", False))

def grant_access(reason="", email=None):
    st.session_state["ad_access_granted"] = True
    st.session_state["ad_access_reason"]  = reason
    if email:
        st.session_state["ad_access_email"] = email

def check_coupon(code):
    raw = get_secret("AD_COUPON_CODES", get_secret("COUPON_CODES", ""))
    valid = {c.strip().upper() for c in raw.split(",") if c.strip()}
    return code.strip().upper() in valid

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
  html, body, [class*="css"] {{ font-family: 'Segoe UI', Arial, sans-serif; }}
  .block-container {{ max-width: 760px; padding-top: 1.5rem; }}
  .header-block {{ background:{PRIMARY}; padding:2rem 2.5rem 1.6rem; border-radius:12px; margin-bottom:2rem; border-bottom:3px solid {GOLD}; }}
  .header-block .logo-wrap {{ text-align:center; margin-bottom:1.1rem; }}
  .header-block .logo-wrap img {{ height:140px; max-width:100%; object-fit:contain; }}
  .header-block .eyebrow {{ color:{GOLD}; font-size:0.72rem; font-weight:700; letter-spacing:0.18em; text-transform:uppercase; margin:0 0 0.3rem; text-align:center; }}
  .header-block h1 {{ color:#fff; font-size:2rem; font-weight:900; margin:0 0 0.35rem; line-height:1.15; text-align:center; }}
  .header-block .sub {{ color:#aaa; font-size:0.95rem; margin:0; text-align:center; }}

  .section-h2 {{ color:{PRIMARY}; font-size:1.3rem; font-weight:900; border-bottom:3px solid {GOLD}; padding-bottom:6px; margin:2rem 0 1rem; }}
  .section-label {{ color:{NAVY}; font-size:0.75rem; font-weight:700; letter-spacing:0.14em; text-transform:uppercase; margin-bottom:0.4rem; }}

  /* Best pick badge */
  .best-badge {{ background:{GOLD}; color:{PRIMARY}; font-weight:700; font-size:0.72rem; padding:2px 10px; border-radius:20px; display:inline-block; margin-bottom:0.35rem; letter-spacing:0.05em; text-transform:uppercase; }}

  /* Headlines */
  .hl-type {{ font-size:0.72rem; font-weight:700; color:{GOLD_DARK}; text-transform:uppercase; letter-spacing:0.12em; margin-bottom:0.2rem; }}
  .hl-card {{ background:{GOLD_LIGHT}; border:1px solid #E8D9B0; border-left:4px solid {GOLD}; border-radius:0 8px 8px 0; padding:0.85rem 1.1rem; margin:0.5rem 0; }}
  .hl-card.best {{ border-left:4px solid {GOLD_DARK}; background:{GOLD_MID}; }}
  .hl-text {{ font-weight:700; font-size:1.08rem; color:{PRIMARY}; margin-bottom:0.25rem; }}
  .hl-meta {{ font-size:0.8rem; color:#888; }}
  .hl-runner {{ font-size:0.82rem; color:#999; font-style:italic; margin-top:0.2rem; }}

  /* Primary text */
  .pt-card {{ background:#fff; border:1px solid #e0d9c8; border-radius:8px; padding:1rem 1.3rem; margin:0.75rem 0; }}
  .pt-card.best {{ border:2px solid {GOLD}; }}
  .pt-label {{ font-weight:700; font-size:0.82rem; color:{GOLD_DARK}; text-transform:uppercase; letter-spacing:0.1em; margin-bottom:0.5rem; }}
  .pt-teaser {{ font-size:0.88rem; color:{NAVY}; font-style:italic; margin-bottom:0.5rem; border-left:3px solid {NAVY}; padding-left:0.6rem; }}
  .pt-excerpt {{ font-size:0.92rem; line-height:1.75; color:#333; white-space:pre-wrap; background:#FAFAFA; border-radius:6px; padding:0.75rem; margin:0.5rem 0; font-family:'Georgia',serif; }}
  .pt-cta {{ font-size:0.88rem; color:{GOLD_DARK}; font-weight:600; margin-top:0.5rem; }}
  .pt-notes {{ font-size:0.8rem; color:#888; margin-top:0.5rem; font-style:italic; }}

  /* Description */
  .desc-card {{ background:{GOLD_BG}; border-left:4px solid {GOLD}; border-radius:0 8px 8px 0; padding:0.75rem 1rem; margin:0.5rem 0; }}
  .desc-card.best {{ border-left:4px solid {GOLD_DARK}; }}
  .desc-label {{ font-size:0.72rem; font-weight:700; color:{GOLD_DARK}; text-transform:uppercase; letter-spacing:0.1em; margin-bottom:0.3rem; }}
  .desc-text {{ font-size:0.95rem; color:{PRIMARY}; font-weight:500; margin-bottom:0.2rem; line-height:1.6; }}
  .desc-chars {{ font-size:0.78rem; color:#aaa; }}

  /* Display URL / CTA */
  .url-card {{ background:{NAVY_BG}; border-left:4px solid {NAVY}; border-radius:0 8px 8px 0; padding:0.9rem 1.1rem; margin:0.5rem 0; }}
  .url-card .url-text {{ font-weight:700; font-size:1rem; color:{NAVY}; font-family:'Courier New', monospace; }}
  .url-card .url-note {{ font-size:0.82rem; color:#555; margin-top:0.4rem; line-height:1.5; }}
  .cta-card {{ background:{NAVY_LIGHT}; border-left:4px solid {NAVY}; border-radius:0 8px 8px 0; padding:0.75rem 1rem; margin:0.5rem 0; }}
  .cta-card .cta-btn {{ font-weight:900; font-size:1.1rem; color:{NAVY}; }}
  .cta-card .cta-reason {{ font-size:0.85rem; color:#444; margin-top:0.3rem; }}

  /* Targeting */
  .targeting-box {{ background:#EEF4FF; border:1px solid #c5d8f8; border-left:4px solid #3B6FD4; border-radius:0 8px 8px 0; padding:0.85rem 1.1rem; margin:0.5rem 0; font-size:0.92rem; color:#1a3a6b; line-height:1.6; }}

  /* Setup reminder */
  .setup-box {{ background:{PRIMARY}; border-radius:10px; padding:1.2rem 1.5rem; margin:1.5rem 0; border-bottom:3px solid {GOLD}; }}
  .setup-box .setup-title {{ color:{GOLD}; font-weight:700; font-size:0.85rem; text-transform:uppercase; letter-spacing:0.12em; margin-bottom:0.6rem; }}
  .setup-box .setup-item {{ color:#ccc; font-size:0.87rem; padding:3px 0; line-height:1.5; }}
  .setup-box .setup-item::before {{ content:"✓ "; color:{GOLD}; font-weight:700; }}

  /* Upgrade */
  .upgrade-card {{ background:#fff; border:2px solid {GOLD_MID}; border-top:4px solid {GOLD}; border-radius:14px; padding:2rem; margin:1.5rem 0; text-align:center; }}
  .upgrade-card .uc-eyebrow {{ font-size:0.75rem; font-weight:700; letter-spacing:0.15em; text-transform:uppercase; color:{GOLD_DARK}; margin-bottom:0.4rem; }}
  .upgrade-card h3 {{ font-size:1.4rem; font-weight:900; color:{PRIMARY}; margin:0 0 0.3rem; }}
  .upgrade-card .uc-sub {{ color:#666; font-size:0.92rem; margin-bottom:1.2rem; }}
  .upgrade-card ul {{ text-align:left; display:inline-block; list-style:none; padding:0; margin:0 auto 1.5rem; }}
  .upgrade-card ul li {{ padding:3px 0; font-size:0.92rem; color:#333; }}
  .upgrade-card ul li::before {{ content:"✓ "; color:{GOLD_DARK}; font-weight:700; }}
  .price-badge {{ display:inline-block; background:{PRIMARY}; color:{GOLD}; font-size:1.6rem; font-weight:900; padding:6px 28px; border-radius:40px; margin-bottom:1.2rem; border:2px solid {GOLD}; }}
  .divider-or {{ display:flex; align-items:center; gap:12px; color:#bbb; font-size:0.85rem; margin:1.2rem 0; }}
  .divider-or::before, .divider-or::after {{ content:""; flex:1; border-top:1px solid #e0e0e0; }}
  .coupon-section {{ background:{GOLD_BG}; border:1px solid #e8d78a; border-left:4px solid {GOLD}; border-radius:0 8px 8px 0; padding:1rem 1.25rem; margin-bottom:1rem; }}
  .success-banner {{ background:{GREEN_BG}; border:1px solid #b8ddc8; border-left:4px solid {GREEN}; border-radius:0 8px 8px 0; padding:1rem 1.25rem; margin-bottom:1.5rem; font-size:0.92rem; color:#155724; }}
  .bookmark-tip {{ background:#EEF4FF; border:1px solid #c5d8f8; border-left:4px solid #3B6FD4; border-radius:0 8px 8px 0; padding:0.85rem 1.1rem; margin-bottom:1.5rem; font-size:0.88rem; color:#1a3a6b; }}
  .stButton > button {{ background:{PRIMARY} !important; color:{GOLD} !important; font-weight:700 !important; padding:0.65rem 2rem !important; border-radius:8px !important; border:2px solid {GOLD} !important; font-size:1rem !important; width:100%; }}
  .stButton > button:hover {{ background:{GOLD} !important; color:{PRIMARY} !important; }}
  .stLinkButton > a {{ background:{PRIMARY} !important; color:{GOLD} !important; font-weight:700 !important; border-radius:8px !important; font-size:1rem !important; border:2px solid {GOLD} !important; }}
  .footer-note {{ text-align:center; color:#aaa; font-size:0.78rem; margin-top:3rem; padding-top:1rem; border-top:1px solid {GOLD_MID}; }}
</style>
""", unsafe_allow_html=True)

# ── AI Prompt ─────────────────────────────────────────────────────────────────
AD_SYSTEM_PROMPT = """You are the Facebook Ad Copy Generator for The Writing Wives Genre & Trope AI Classroom.

Generate a complete, ready-to-paste Facebook & Instagram ad copy package based on the book details provided.

Return ONLY valid JSON. No markdown, no code fences. Exact structure:

{
  "primary_text": {
    "teaser": "Pre-qualifier line (e.g. 'If you love dark romance and morally grey heroes...')",
    "excerpt": "The trimmed excerpt or excerpt-style hook text (300-400 words max, ends on tension or cliffhanger). Use \\n\\n for paragraph breaks.",
    "cta_line": "Short CTA line (e.g. 'Available now — read free in Kindle Unlimited' or 'Get your copy today →')"
  },
  "headlines": [
    {"type": "Trope Declaration", "text": "Enemies to lovers. Small town. All the heat.", "chars": 42, "runner_up": "Grumpy hero. Sunshine heroine. Slow burn."},
    {"type": "CTA with Hook",     "text": "Read the first chapter free today",           "chars": 34, "runner_up": "Start the series that readers can't put down"},
    {"type": "Reader Identity",   "text": "For readers who like their heroes morally grey","chars": 47, "runner_up": "If you love slow burns that earn every page"},
    {"type": "Series Binge Hook", "text": "Book 1 of 5. Start the binge today.",          "chars": 36, "runner_up": "5 books. One obsession. Start here."},
    {"type": "Genre Positioning", "text": "Dark paranormal romance. High heat. HEA.",     "chars": 42, "runner_up": "Slow-burn fantasy romance. Fated mates."}
  ],
  "description_short": {
    "text": "Under 200 characters. Dense trope and genre keywords. Sounds like ad copy, not a keyword list.",
    "chars": 145
  },
  "description_long": {
    "text": "Up to 500 characters. GEO-style: trope declaration + comp title callout + optional reader quote. Front-load the most searchable trope terms.",
    "chars": 380
  },
  "display_url": "amazon.com/kindle-unlimited",
  "cta_recommendation": {
    "button": "Download",
    "reasoning": "One sentence explaining why this CTA button is best for this book and audience."
  },
  "targeting_note": "1-2 sentences on Facebook audience targeting: who to target, what interest stacks work for this genre.",
  "best_headline": 0
}

Rules:
- Headlines: Under 40 characters preferred (flag if over). Never open with the book title. Lead with trope/hook/tension. Use genre shorthand readers know: HEA, HFN, KU, MM, RH, MFM.
- Primary text: First word NOT 'I'. Open with the teaser pre-qualifier, then the excerpt or excerpt-style hook, then the CTA line.
- If no excerpt was provided, write compelling excerpt-STYLE primary text from the genre/tropes — but keep it feeling like a scene, not ad copy.
- description_short: 150-200 chars. One clear reader promise or trope signal.
- description_long: Up to 500 chars. GEO-style — trope declaration, comp title, optionally a trope-specific reader promise.
- display_url: Use 'amazon.com/kindle-unlimited' for KU, 'amazon.com/books' for purchase.
- cta_recommendation button: 'Download' (best for most book ads), 'Shop Now' (print/direct), 'Learn More' (soft sell only).
- best_headline: index (0-4) of the strongest headline for cold traffic."""

def call_openrouter_ads(book_title, genre, tropes, blurb, excerpt, ku_or_purchase, series_or_standalone):
    api_key = get_secret("OPENROUTER_API_KEY")
    if not api_key:
        st.error("API key not configured.")
        st.stop()
    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
    model  = get_secret("MODEL", "anthropic/claude-sonnet-4-5")
    parts  = []
    if book_title:          parts.append(f"Book title: {book_title}")
    if genre:               parts.append(f"Genre: {genre}")
    if tropes:              parts.append(f"Key tropes: {tropes}")
    if ku_or_purchase:      parts.append(f"Availability: {ku_or_purchase}")
    if series_or_standalone:parts.append(f"Series or standalone: {series_or_standalone}")
    if blurb:               parts.append(f"\nBlurb:\n{blurb}")
    if excerpt:             parts.append(f"\nExcerpt (use as primary text basis):\n{excerpt}")
    else:                   parts.append("\nNo excerpt provided — generate excerpt-style primary text from genre/tropes.")
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": AD_SYSTEM_PROMPT},
            {"role": "user",   "content": "\n".join(parts)},
        ],
        temperature=0.5,
        max_tokens=4000,
    )
    raw = response.choices[0].message.content.strip()
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    return json.loads(raw.strip())

# ── Render ────────────────────────────────────────────────────────────────────
def render_ad_package(data, book_title):
    best_hl = data.get("best_headline", 0)

    # ── 1. Primary Text ───────────────────────────────────────────────────────
    st.markdown('<div class="section-h2">📝 Primary Text</div>', unsafe_allow_html=True)
    st.caption("Goes in the Primary Text field in Ads Manager. Teaser → Excerpt → CTA.")
    pt = data.get("primary_text", {})
    if pt:
        teaser  = pt.get("teaser", "")
        excerpt = pt.get("excerpt", "").replace("\\n", "\n")
        cta     = pt.get("cta_line", "")
        st.markdown(f"""
        <div class="pt-card best">
          <div class="pt-label">Primary Text — Copy into Ads Manager</div>
          {"<div class='pt-teaser'>" + teaser + "</div>" if teaser else ""}
          <div class="pt-excerpt">{excerpt}</div>
          {"<div class='pt-cta'>" + cta + "</div>" if cta else ""}
        </div>""", unsafe_allow_html=True)

    # ── 2. Headlines ──────────────────────────────────────────────────────────
    st.markdown('<div class="section-h2">📰 Headlines</div>', unsafe_allow_html=True)
    st.caption("5 typed headlines — one of each angle. Copy all 5 into the Headlines field in Ads Manager and let Facebook test them.")
    headlines = data.get("headlines", [])
    for i, hl in enumerate(headlines):
        is_best = (i == best_hl)
        badge   = '<div class="best-badge">⭐ Best for Cold Traffic</div>' if is_best else ""
        cls     = "hl-card best" if is_best else "hl-card"
        runner  = f'<div class="hl-runner">Runner-up: {hl["runner_up"]}</div>' if hl.get("runner_up") else ""
        chars   = hl.get("chars", len(hl.get("text", "")))
        char_warning = " ⚠️ over 40 chars" if chars > 40 else ""
        st.markdown(f"""
        <div class="{cls}">
          {badge}
          <div class="hl-type">{hl.get("type","")}</div>
          <div class="hl-text">{hl.get("text","")}</div>
          <div class="hl-meta">{chars} chars{char_warning}</div>
          {runner}
        </div>""", unsafe_allow_html=True)

    # ── 3. Description Fields ─────────────────────────────────────────────────
    st.markdown('<div class="section-h2">🏷️ Description Field</div>', unsafe_allow_html=True)
    st.caption("Two versions — use the one that fits the space available in your ad format.")

    desc_short = data.get("description_short", {})
    if desc_short:
        st.markdown(f"""
        <div class="desc-card best">
          <div class="desc-label">Short Version &nbsp;·&nbsp; 150–200 chars — best for most formats</div>
          <div class="desc-text">{desc_short.get("text","")}</div>
          <div class="desc-chars">{desc_short.get("chars", len(desc_short.get("text","")))} characters</div>
        </div>""", unsafe_allow_html=True)

    desc_long = data.get("description_long", {})
    if desc_long:
        st.markdown(f"""
        <div class="desc-card">
          <div class="desc-label">Long Version &nbsp;·&nbsp; up to 500 chars — GEO-style with comp titles</div>
          <div class="desc-text">{desc_long.get("text","")}</div>
          <div class="desc-chars">{desc_long.get("chars", len(desc_long.get("text","")))} characters</div>
        </div>""", unsafe_allow_html=True)

    # ── 4. Display URL & CTA ──────────────────────────────────────────────────
    st.markdown('<div class="section-h2">🔗 Display URL &amp; CTA Button</div>', unsafe_allow_html=True)
    st.caption("Display URL is the text shown under your image — not a functional link. Set it in the Display Link field, not the Destination URL.")

    display_url = data.get("display_url", "")
    if display_url:
        st.markdown(f"""
        <div class="url-card">
          <div class="url-text">{display_url}</div>
          <div class="url-note">Enter this in Ads Manager → Ad creative → Display link field.</div>
        </div>""", unsafe_allow_html=True)

    cta_rec = data.get("cta_recommendation", {})
    if cta_rec:
        st.markdown(f"""
        <div class="cta-card">
          <div class="cta-btn">CTA Button: {cta_rec.get("button","")}</div>
          <div class="cta-reason">{cta_rec.get("reasoning","")}</div>
        </div>""", unsafe_allow_html=True)

    # ── 5. Targeting ──────────────────────────────────────────────────────────
    if data.get("targeting_note"):
        st.markdown('<div class="section-h2">🎯 Audience Targeting Suggestion</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="targeting-box">🎯 {data["targeting_note"]}</div>', unsafe_allow_html=True)

    # ── 6. Quick Setup Reminder ───────────────────────────────────────────────
    st.markdown(f"""
    <div class="setup-box">
      <div class="setup-title">⚡ Quick Setup Reminder — Before You Publish</div>
      <div class="setup-item">Turn off Dynamic Creative (it overwrites your tested copy)</div>
      <div class="setup-item">Turn off all AI enhancements except Add Overlays</div>
      <div class="setup-item">Set your Display Link in the creative panel — it's easy to miss</div>
      <div class="setup-item">Enter all 5 headlines so Facebook can test them automatically</div>
      <div class="setup-item">Run through your Module 4/6 pre-launch checklist before publishing</div>
    </div>
    """, unsafe_allow_html=True)

    # ── Download ──────────────────────────────────────────────────────────────
    st.divider()
    st.markdown("**Download your ad package:**")
    ca, cb = st.columns(2)
    with ca:
        st.download_button("⬇ Download as JSON",
            data=json.dumps(data, indent=2, ensure_ascii=False).encode(),
            file_name=f"fb-ad-package-{(book_title or 'report').replace(' ','-').lower()}.json",
            mime="application/json", key="dl_ad_json")
    with cb:
        if DOCX_AVAILABLE:
            docx_bytes = generate_ad_docx(data, book_title)
            if docx_bytes:
                st.download_button("⬇ Download as Word Doc",
                    data=docx_bytes,
                    file_name=f"FB Ad Package - {book_title or 'Report'}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_ad_docx")

# ── Word Doc ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_navy_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1C, 0x3A, 0x5E)
    run.font.name = "Segoe UI"
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "C8A200")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_shaded_box(doc, lines, fill_hex="FFFBEC", border_color="C8A200"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, fill_hex)
    cell._tc.get_or_add_tcPr()
    for tag in cell._tc.iter(qn("w:tcBorders")):
        for side in ["top","bottom","left","right","insideH","insideV"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "12")
            el.set(qn("w:color"), border_color)
            tag.append(el)
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return tbl

def add_label(doc, text, color_rgb=(0xA8, 0x86, 0x3A)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.name = "Segoe UI"
    return p

def generate_ad_docx(data, book_title):
    try:
        doc = Document()
        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        doc.styles["Normal"].font.name = "Segoe UI"
        doc.styles["Normal"].font.size = Pt(11)

        # ── Title block ───────────────────────────────────────────────────────
        title_tbl = doc.add_table(rows=1, cols=1)
        title_tbl.style = "Table Grid"
        tc = title_tbl.rows[0].cells[0]
        set_cell_bg(tc, "1A1A1A")
        for tag in tc._tc.iter(qn("w:tcBorders")):
            for side in ["top","bottom","left","right"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "12")
                el.set(qn("w:color"), "C8A200")
                tag.append(el)
        tc.paragraphs[0].clear()
        ey = tc.paragraphs[0]
        ey.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ey.paragraph_format.space_before = Pt(12)
        er = ey.add_run("GENRE & TROPE AI CLASSROOM")
        er.font.size = Pt(8); er.font.color.rgb = RGBColor(0xAA,0xAA,0xAA); er.font.name = "Segoe UI"
        tp = tc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(4); tp.paragraph_format.space_after = Pt(4)
        tr = tp.add_run("Facebook Ad Copy Package")
        tr.bold = True; tr.font.size = Pt(22); tr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); tr.font.name = "Segoe UI"
        sp = tc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(12)
        sr = sp.add_run(book_title or "Ad Copy Package")
        sr.font.size = Pt(13); sr.font.color.rgb = RGBColor(0xC8,0xA2,0x00); sr.font.name = "Segoe UI"
        doc.add_paragraph()

        # ── Book at a glance ──────────────────────────────────────────────────
        add_navy_heading(doc, "Book at a Glance")
        doc.add_paragraph()

        # ── Primary Text ──────────────────────────────────────────────────────
        add_navy_heading(doc, "1. Primary Text")
        add_label(doc, "→ Copy into Ads Manager: Primary Text field")
        pt = data.get("primary_text", {})
        if pt:
            lines = []
            if pt.get("teaser"):
                lines.append(pt["teaser"])
                lines.append("")
            for para in pt.get("excerpt","").replace("\\n","\n").split("\n"):
                lines.append(para)
            if pt.get("cta_line"):
                lines.append("")
                lines.append(pt["cta_line"])
            add_shaded_box(doc, lines, "FAFAFA", "1C3A5E")
        doc.add_paragraph()

        # ── Headlines ─────────────────────────────────────────────────────────
        add_navy_heading(doc, "2. Headlines")
        add_label(doc, "→ Copy all 5 into Ads Manager: Headlines field")
        best_hl = data.get("best_headline", 0)
        for i, hl in enumerate(data.get("headlines", [])):
            is_best = (i == best_hl)
            add_label(doc, f"TYPE {i+1}: {hl.get('type','').upper()}{' — ⭐ BEST FOR COLD TRAFFIC' if is_best else ''}")
            lines = [hl.get("text",""), f"({hl.get('chars', len(hl.get('text','')))} chars)"]
            if hl.get("runner_up"):
                lines += ["", f"Runner-up: {hl['runner_up']}"]
            add_shaded_box(doc, lines, "FAF4E6" if is_best else "F5F5F5", "C8A200" if is_best else "CCCCCC")
            doc.add_paragraph()

        # ── Description Fields ────────────────────────────────────────────────
        add_navy_heading(doc, "3. Description Field")
        add_label(doc, "→ Copy into Ads Manager: Description field")

        desc_short = data.get("description_short", {})
        if desc_short:
            add_label(doc, f"SHORT VERSION ({desc_short.get('chars', '')} chars) — best for most formats")
            add_shaded_box(doc, [desc_short.get("text","")], "FFFBEC", "C8A200")
            doc.add_paragraph()

        desc_long = data.get("description_long", {})
        if desc_long:
            add_label(doc, f"LONG VERSION ({desc_long.get('chars', '')} chars) — GEO-style with comp titles")
            add_shaded_box(doc, [desc_long.get("text","")], "FAF4E6", "C8A200")
            doc.add_paragraph()

        # ── Display URL & CTA ─────────────────────────────────────────────────
        add_navy_heading(doc, "4. Display URL & CTA Button")
        add_label(doc, "→ Display URL goes in Ads Manager: Display Link field (not the Destination URL)")
        if data.get("display_url"):
            add_shaded_box(doc, [data["display_url"]], "DDE8F4", "1C3A5E")
        cta_rec = data.get("cta_recommendation", {})
        if cta_rec:
            add_label(doc, f"CTA BUTTON: {cta_rec.get('button','').upper()}")
            p = doc.add_paragraph(cta_rec.get("reasoning",""))
            p.paragraph_format.space_after = Pt(8)
        doc.add_paragraph()

        # ── Targeting ─────────────────────────────────────────────────────────
        if data.get("targeting_note"):
            add_navy_heading(doc, "5. Audience Targeting Suggestion")
            p = doc.add_paragraph(data["targeting_note"])
            p.paragraph_format.space_after = Pt(8)
            doc.add_paragraph()

        # ── Setup Reminder ────────────────────────────────────────────────────
        reminder_tbl = doc.add_table(rows=1, cols=1)
        reminder_tbl.style = "Table Grid"
        rc = reminder_tbl.rows[0].cells[0]
        set_cell_bg(rc, "1A1A1A")
        for tag in rc._tc.iter(qn("w:tcBorders")):
            for side in ["top","bottom","left","right"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "12"); el.set(qn("w:color"), "C8A200")
                tag.append(el)
        rc.paragraphs[0].clear()
        hp = rc.paragraphs[0]
        hp.paragraph_format.space_before = Pt(8); hp.paragraph_format.space_after = Pt(6)
        hr = hp.add_run("⚡  QUICK SETUP REMINDER — BEFORE YOU PUBLISH")
        hr.bold = True; hr.font.size = Pt(9); hr.font.color.rgb = RGBColor(0xC8,0xA2,0x00); hr.font.name = "Segoe UI"
        reminders = [
            "Turn off Dynamic Creative (it overwrites your tested copy)",
            "Turn off all AI enhancements except Add Overlays",
            "Set your Display Link in the creative panel — it's easy to miss",
            "Enter all 5 headlines so Facebook can test them automatically",
            "Run through your Module 4/6 pre-launch checklist before publishing",
        ]
        for item in reminders:
            rp = rc.add_paragraph()
            rp.paragraph_format.space_before = Pt(2); rp.paragraph_format.space_after = Pt(2)
            rr = rp.add_run(f"✓  {item}")
            rr.font.size = Pt(9.5); rr.font.color.rgb = RGBColor(0xCC,0xCC,0xCC); rr.font.name = "Segoe UI"
        ep = rc.add_paragraph(); ep.paragraph_format.space_before = Pt(4)

        # ── Footer ────────────────────────────────────────────────────────────
        doc.add_paragraph()
        fp = doc.add_paragraph("The Writing Wives  ·  Genre & Trope AI Classroom  ·  thewritingwives.com")
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.runs[0].font.size = Pt(8)
        fp.runs[0].font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        st.error(f"Word doc generation error: {e}")
        return None

def show_upgrade_card():
    price = get_secret("AD_PRICE_DISPLAY", "$25")
    sub   = get_secret("AD_PRICE_SUBTITLE", "per report · ad copy for your blurb")
    link  = get_secret("STRIPE_AD_LINK", "")
    st.markdown(f"""
    <div class="upgrade-card">
      <div class="uc-eyebrow">Unlock Your Ad Package</div>
      <h3>Ready-to-paste copy for every ad element</h3>
      <p class="uc-sub">Paste your blurb and excerpt — get a complete Facebook &amp; Instagram ad package.</p>
      <ul>
        <li>Primary text — teaser + excerpt + CTA, formatted for mobile</li>
        <li>5 typed headlines — one of each angle, under 40 chars</li>
        <li>Short &amp; long description field options</li>
        <li>Display URL + CTA button recommendation</li>
        <li>⭐ Best Pick for cold traffic highlighted</li>
        <li>Audience targeting suggestion</li>
        <li>Quick Setup Reminder checklist</li>
        <li>Downloadable Word doc — open and copy straight to Ads Manager</li>
      </ul>
      <div class="price-badge">{price}</div>
      <div style="color:#888;font-size:0.85rem;margin-bottom:1.2rem;">{sub}</div>
    </div>
    """, unsafe_allow_html=True)
    if link:
        st.link_button(f"Get My Ad Package — {price} →", link, use_container_width=True)
    st.markdown('<div class="divider-or">or</div>', unsafe_allow_html=True)
    st.markdown('<div class="coupon-section"><div class="label">🎓 Writing Wives Skool Member? Enter your coupon for free access.</div></div>', unsafe_allow_html=True)
    with st.form("ad_coupon_form"):
        code  = st.text_input("Coupon code", placeholder="e.g. WRITINGWIVES", label_visibility="collapsed")
        apply = st.form_submit_button("Apply Coupon →")
    if apply:
        if code and check_coupon(code):
            grant_access(reason="coupon")
            st.rerun()
        else:
            st.error("That coupon code isn't valid.")

# ── Stripe URL check ──────────────────────────────────────────────────────────
if not is_authenticated():
    sid = st.query_params.get("ad_session_id")
    if sid and not st.session_state.get("ad_stripe_checked"):
        st.session_state["ad_stripe_checked"] = True
        with st.spinner("Verifying your payment..."):
            ok, email = verify_stripe_session(sid)
        if ok:
            grant_access(reason="stripe", email=email)
            st.rerun()
        else:
            st.warning("Payment couldn't be verified yet — try refreshing in a moment.")

# ── Header ────────────────────────────────────────────────────────────────────
if st.button("← Back to All Tools", key="back_home"):
    st.switch_page("home.py")

_logo_b64 = get_logo_b64()
_logo_html = f'<div class="logo-wrap"><img src="data:image/png;base64,{_logo_b64}" alt="The Writing Wives"></div>' if _logo_b64 else ""
st.markdown(f"""
<div class="header-block">
  {_logo_html}
  <div class="eyebrow">Genre &amp; Trope AI Classroom</div>
  <h1>FB &amp; Instagram Ad Package</h1>
  <p class="sub">Paste your blurb and a chapter excerpt — get a complete, copy-paste-ready ad package for Ads Manager.</p>
</div>
""", unsafe_allow_html=True)

# ── Access banners ────────────────────────────────────────────────────────────
if is_authenticated():
    reason = st.session_state.get("ad_access_reason", "")
    email  = st.session_state.get("ad_access_email", "")
    if reason == "stripe":
        st.markdown(f'<div class="success-banner">✅ <strong>Payment confirmed{f" for {email}" if email else ""}.</strong> Full access unlocked.</div>', unsafe_allow_html=True)
        st.markdown('<div class="bookmark-tip">🔖 <strong>Bookmark this page</strong> — your URL contains your access token.</div>', unsafe_allow_html=True)
    elif reason in ("coupon", "lifetime"):
        st.markdown('<div class="success-banner">🎓 <strong>Access granted.</strong> Generate as many packages as you need.</div>', unsafe_allow_html=True)

# ── Form ──────────────────────────────────────────────────────────────────────
with st.form("ad_form"):
    book_title = st.text_input("Book Title", placeholder="e.g. We Are Legion (We Are Bob)")

    col1, col2 = st.columns(2)
    with col1:
        genre  = st.text_input("Genre / Sub-Genre", placeholder="e.g. Dark Paranormal Romance")
        ku_or_purchase = st.selectbox("Availability", ["Kindle Unlimited (KU)", "Purchase only", "Both KU and purchase"])
    with col2:
        tropes = st.text_input("Key Tropes", placeholder="e.g. Enemies to Lovers, Forced Proximity")
        series_or_standalone = st.selectbox("Series or Standalone?", ["Series — Book 1", "Series — Mid-series", "Standalone"])

    blurb = st.text_area("Your Blurb", height=140,
        placeholder="Paste your Amazon blurb here...")

    excerpt = st.text_area("Chapter 1 Excerpt (optional but recommended)",
        height=200,
        placeholder="Paste 200–500 words from your first chapter or a strong hook scene. "
                     "This becomes your primary text. If left blank, the AI will write excerpt-style copy from your blurb and tropes.")

    submit_label = "Generate My Ad Package →" if is_authenticated() else "See What's Included →"
    submitted = st.form_submit_button(submit_label)

# ── On submit ─────────────────────────────────────────────────────────────────
if submitted:
    if not blurb.strip() and not excerpt.strip():
        st.warning("Please paste at least your blurb or an excerpt to generate ad copy.")
    elif not is_authenticated():
        # Teaser
        st.info("Here's everything your ad package includes — unlock to generate yours.")
        st.markdown(f"""
        <div style="background:{GOLD_LIGHT};border:1px solid #E8D9B0;border-radius:10px;padding:1.2rem 1.5rem;margin:1rem 0;line-height:2;">
          <strong>📝 Primary Text</strong> — Teaser pre-qualifier + your excerpt formatted for mobile scroll + CTA line<br>
          <strong>📰 5 Typed Headlines</strong> — Trope Declaration · CTA Hook · Reader Identity · Binge Hook · Genre Positioning<br>
          <strong>🏷️ Short &amp; Long Description</strong> — 150-char and 500-char versions, GEO-optimised<br>
          <strong>🔗 Display URL + CTA Button</strong> — Exact text for the Display Link field + button recommendation<br>
          <strong>🎯 Targeting Suggestion</strong> — Audience interest stacks for your genre<br>
          <strong>⚡ Setup Reminder</strong> — Pre-launch checklist so nothing gets missed in Ads Manager<br>
          <strong>⬇ Word Doc</strong> — Open alongside Ads Manager and copy-paste directly
        </div>
        """, unsafe_allow_html=True)
        show_upgrade_card()
    else:
        with st.spinner("Writing your ad copy..."):
            try:
                result = call_openrouter_ads(
                    book_title, genre, tropes, blurb,
                    excerpt, ku_or_purchase, series_or_standalone
                )
                st.session_state["last_ad_result"] = {"result": result, "book_title": book_title}
            except json.JSONDecodeError:
                st.error("The AI returned an unexpected format. Please try again.")
                st.stop()
            except Exception as e:
                st.error("Something went wrong. Check your API key and try again.")
                st.exception(e)
                st.stop()

        st.success("Your ad package is ready!")
        st.divider()
        render_ad_package(result, book_title)

# ── Coupon unlock with cached result ─────────────────────────────────────────
elif is_authenticated() and st.session_state.get("last_ad_result") and st.session_state.get("ad_access_reason") == "coupon":
    prior = st.session_state["last_ad_result"]
    st.success("✅ Access granted! Here's your ad package:")
    st.divider()
    render_ad_package(prior["result"], prior["book_title"])
    del st.session_state["last_ad_result"]

# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="footer-note">
  The Writing Wives · Genre &amp; Trope AI Classroom ·
  <a href="https://thewritingwives.com" style="color:#999;">thewritingwives.com</a>
</div>
""", unsafe_allow_html=True)
