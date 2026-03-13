import streamlit as st
from openai import OpenAI
import json, re, io, base64
from pathlib import Path

from affiliate_utils import capture_ref, affiliate_payment_link
capture_ref()

try:
    import requests
    from bs4 import BeautifulSoup
    SCRAPING_AVAILABLE = True
except ImportError:
    SCRAPING_AVAILABLE = False

try:
    import stripe as stripe_lib
    STRIPE_AVAILABLE = True
except ImportError:
    STRIPE_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False



PRIMARY    = "#1A1A1A"
PRIMARY_MID= "#2E2E2E"
GOLD       = "#D4B36E"
GOLD_DARK  = "#A8863A"
GOLD_LIGHT = "#FAF4E6"
GOLD_MID   = "#F5EDD6"
GOLD_BG    = "#FFFBEC"
GREEN      = "#1A7A4A"
GREEN_BG   = "#D4EDDA"

CATEGORY_LABELS = {
    "title_subtitle":      "1. Title & Subtitle",
    "blurb_effectiveness": "2. Blurb Effectiveness",
    "review_health":       "3. Review Health",
    "price_positioning":   "4. Price Positioning",
    "discoverability":     "5. Discoverability",
    "conversion_readiness":"6. Conversion Readiness",
}

SCRAPE_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "en-US,en;q=0.9",
    "Accept": (
        "text/html,application/xhtml+xml,application/xml;"
        "q=0.9,image/webp,*/*;q=0.8"
    ),
    "Accept-Encoding": "gzip, deflate, br",
    "Connection": "keep-alive",
    "Upgrade-Insecure-Requests": "1",
    "Cache-Control": "max-age=0",
}

# ── Helpers ───────────────────────────────────────────────────────────────────
def get_logo_b64():
    p = Path(__file__).parent / "logo.png"
    if p.exists():
        with open(p, "rb") as f:
            return base64.b64encode(f.read()).decode()
    return None

def get_secret(key, default=""):
    try:
        return st.secrets.get(key, default)
    except Exception:
        return default

def verify_stripe_session(session_id):
    if not STRIPE_AVAILABLE:
        return False, None
    try:
        stripe_lib.api_key = get_secret("STRIPE_SECRET_KEY")
        session = stripe_lib.checkout.Session.retrieve(session_id)
        if session.payment_status == "paid":
            email = session.customer_details.email if session.customer_details else None
            return True, email
    except Exception:
        pass
    return False, None

def is_authenticated():
    return (st.session_state.get("amazon_access_granted", False) or
            st.session_state.get("lifetime_access", False))

def grant_access(reason="", email=None):
    st.session_state["amazon_access_granted"] = True
    st.session_state["amazon_access_reason"]  = reason
    if email:
        st.session_state["amazon_access_email"] = email

def check_coupon(code):
    raw = get_secret("AMAZON_COUPON_CODES", get_secret("COUPON_CODES", ""))
    valid = {c.strip().upper() for c in raw.split(",") if c.strip()}
    return code.strip().upper() in valid

# ── Amazon scraping ───────────────────────────────────────────────────────────
def extract_asin(url: str):
    """Pull the 10-character ASIN out of any standard Amazon URL."""
    for pattern in [r"/dp/([A-Z0-9]{10})", r"/gp/product/([A-Z0-9]{10})"]:
        m = re.search(pattern, url)
        if m:
            return m.group(1)
    return None

def clean_url(url: str) -> str:
    """Normalise to a clean dp URL."""
    asin = extract_asin(url)
    if asin:
        return f"https://www.amazon.com/dp/{asin}"
    return url.strip()

def scrape_amazon_page(url: str) -> dict:
    """
    Fetch an Amazon product page and extract key fields.
    Returns a dict with keys: title, subtitle, author, blurb,
    star_rating, review_count, price, categories, asin, success, error.
    """
    result = {
        "title": "", "subtitle": "", "author": "", "blurb": "",
        "star_rating": None, "review_count": None, "price": "",
        "categories": "", "asin": extract_asin(url) or "",
        "success": False, "error": "",
    }

    if not SCRAPING_AVAILABLE:
        result["error"] = "Scraping libraries not installed."
        return result

    url = clean_url(url)

    try:
        resp = requests.get(url, headers=SCRAPE_HEADERS, timeout=15)
        if resp.status_code != 200:
            result["error"] = (
                f"Amazon returned a {resp.status_code} response. "
                "This usually means the request was blocked — please fill in the details manually below."
            )
            return result
    except requests.RequestException as e:
        result["error"] = f"Could not reach Amazon: {e}"
        return result

    soup = BeautifulSoup(resp.text, "lxml")

    # ── Title ──────────────────────────────────────────────────────────────────
    title_el = soup.select_one("#productTitle")
    if title_el:
        result["title"] = title_el.get_text(strip=True)

    # ── Subtitle ───────────────────────────────────────────────────────────────
    subtitle_el = soup.select_one("#productSubtitle")
    if subtitle_el:
        result["subtitle"] = subtitle_el.get_text(strip=True)

    # ── Author ─────────────────────────────────────────────────────────────────
    author_el = soup.select_one("#bylineInfo .author .a-link-normal, #bylineInfo .contributorNameID")
    if not author_el:
        author_el = soup.select_one(".author .a-link-normal")
    if author_el:
        result["author"] = author_el.get_text(strip=True)

    # ── Star rating ────────────────────────────────────────────────────────────
    rating_el = soup.select_one("span.a-icon-alt")
    if rating_el:
        m = re.search(r"([\d.]+) out of 5", rating_el.get_text())
        if m:
            result["star_rating"] = float(m.group(1))

    if not result["star_rating"]:
        rating_el2 = soup.select_one("#acrPopover")
        if rating_el2 and rating_el2.get("title"):
            m = re.search(r"([\d.]+) out of 5", rating_el2["title"])
            if m:
                result["star_rating"] = float(m.group(1))

    # ── Review count ───────────────────────────────────────────────────────────
    review_el = soup.select_one("#acrCustomerReviewText")
    if review_el:
        m = re.search(r"([\d,]+)", review_el.get_text())
        if m:
            result["review_count"] = int(m.group(1).replace(",", ""))

    # ── Price ──────────────────────────────────────────────────────────────────
    for selector in [
        "#kindle-price .a-offscreen",
        ".kindle-price .a-offscreen",
        "#price .a-offscreen",
        ".a-price .a-offscreen",
        "#listPrice",
        "#buyNewSection .a-offscreen",
    ]:
        price_el = soup.select_one(selector)
        if price_el:
            price_text = price_el.get_text(strip=True)
            if price_text and "$" in price_text:
                result["price"] = price_text
                break

    # ── Blurb / description ────────────────────────────────────────────────────
    desc_el = soup.select_one("#bookDescription_feature_div")
    if not desc_el:
        desc_el = soup.select_one("#productDescription")
    if desc_el:
        # Remove "Read more" / "Read less" toggle spans
        for tag in desc_el.select("span.a-expander-content, noscript, .a-expander-header"):
            tag.decompose()
        raw_desc = desc_el.get_text(separator="\n", strip=True)
        # Clean up excess whitespace
        lines = [l.strip() for l in raw_desc.splitlines() if l.strip()]
        result["blurb"] = "\n\n".join(lines)

    # ── Categories ─────────────────────────────────────────────────────────────
    cats = []
    for el in soup.select("#wayfinding-breadcrumbs_feature_div li, .zg_hrsr .zg_hrsr_ladder"):
        text = el.get_text(strip=True)
        if text and text not in ("›", "»", ""):
            cats.append(text)
    if cats:
        result["categories"] = " > ".join(cats[:6])

    # ── Bot-detection check ────────────────────────────────────────────────────
    page_text = soup.get_text()
    if (not result["title"] and
            ("Enter the characters you see below" in page_text or
             "Type the characters you see in this image" in page_text or
             "Robot Check" in page_text)):
        result["error"] = (
            "Amazon served a CAPTCHA / bot-detection page. "
            "Please fill in the details manually below."
        )
        return result

    if not result["title"] and not result["blurb"]:
        result["error"] = (
            "Couldn't extract data from that page — "
            "Amazon may have blocked the request. "
            "Please fill in the details manually below."
        )
        return result

    result["success"] = True
    return result

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
  html, body, [class*="css"] {{ font-family: 'Segoe UI', Arial, sans-serif; }}
  .block-container {{ max-width: 760px; padding-top: 1.5rem; }}
  .header-block {{ background:{PRIMARY}; padding:2rem 2.5rem 1.6rem; border-radius:12px; margin-bottom:2rem; border-bottom:3px solid {GOLD}; }}
  .header-block .logo-wrap {{ text-align:center; margin-bottom:1.1rem; }}
  .header-block .logo-wrap img {{ height:140px; max-width:100%; object-fit:contain; }}
  .header-block .eyebrow {{ color:{GOLD}; font-size:0.72rem; font-weight:700; letter-spacing:0.18em; text-transform:uppercase; margin:0 0 0.3rem; text-align:center; }}
  .header-block h1 {{ color:#fff; font-size:2rem; font-weight:900; margin:0 0 0.35rem; line-height:1.15; text-align:center; }}
  .header-block .sub {{ color:#aaa; font-size:0.95rem; margin:0; text-align:center; }}
  .r-strong  {{ background:#D4EDDA; color:#155724; padding:3px 12px; border-radius:20px; font-weight:700; font-size:0.88rem; display:inline-block; }}
  .r-warning {{ background:#FFF3CD; color:#856404; padding:3px 12px; border-radius:20px; font-weight:700; font-size:0.88rem; display:inline-block; }}
  .r-missing {{ background:#F8D7DA; color:#721C24; padding:3px 12px; border-radius:20px; font-weight:700; font-size:0.88rem; display:inline-block; }}
  .section-h2 {{ color:{PRIMARY}; font-size:1.3rem; font-weight:900; border-bottom:3px solid {GOLD}; padding-bottom:6px; margin:2rem 0 1rem; }}
  .cat-heading {{ color:{PRIMARY}; font-size:1.1rem; font-weight:700; border-bottom:2px solid {GOLD}; padding-bottom:4px; margin:1.6rem 0 0.6rem; }}
  .fix-box {{ background:{GOLD_LIGHT}; border-left:4px solid {GOLD}; padding:0.75rem 1rem; border-radius:0 8px 8px 0; margin-top:0.5rem; font-size:0.92rem; line-height:1.65; }}
  .priority-box {{ background:{GOLD_BG}; border-left:5px solid {GOLD_DARK}; padding:0.9rem 1.1rem; border-radius:0 10px 10px 0; margin:0.4rem 0; font-size:0.93rem; line-height:1.6; }}
  .sum-grid {{ display:grid; grid-template-columns:1fr 1fr; gap:8px; margin:0.75rem 0 1.5rem; }}
  .sum-cell {{ background:{GOLD_MID}; border-radius:8px; padding:10px 14px; display:flex; align-items:center; justify-content:space-between; font-size:0.9rem; border:1px solid #E8D9B0; }}
  .sum-cell .sum-label {{ font-weight:600; color:#333; }}
  .score-big {{ text-align:center; padding:1.5rem; background:{PRIMARY}; border-radius:12px; margin:1rem 0; border-bottom:3px solid {GOLD}; }}
  .score-big .score-num {{ font-size:3.5rem; font-weight:900; color:{GOLD}; line-height:1; }}
  .score-big .score-label {{ color:#aaa; font-size:0.9rem; margin-top:0.3rem; }}
  .score-big .score-status {{ font-size:1rem; font-weight:700; margin-top:0.5rem; }}
  .star-box-good {{ background:{GREEN_BG}; border-left:4px solid {GREEN}; border-radius:0 8px 8px 0; padding:0.85rem 1.1rem; margin:0.5rem 0; font-size:0.93rem; color:#155724; }}
  .star-box-warn {{ background:#FFF3CD; border-left:4px solid #F6A623; border-radius:0 8px 8px 0; padding:0.85rem 1.1rem; margin:0.5rem 0; font-size:0.93rem; color:#856404; }}
  .star-box-bad  {{ background:#F8D7DA; border-left:4px solid #E74C3C; border-radius:0 8px 8px 0; padding:0.85rem 1.1rem; margin:0.5rem 0; font-size:0.93rem; color:#721C24; }}
  .keyword-tag {{ display:inline-block; background:{PRIMARY}; color:{GOLD}; font-size:0.82rem; font-weight:700; padding:4px 12px; border-radius:20px; margin:3px 4px; border:1px solid {GOLD}; }}
  .category-item {{ background:{GOLD_LIGHT}; border-left:3px solid {GOLD}; padding:6px 12px; margin:4px 0; font-size:0.9rem; border-radius:0 6px 6px 0; }}
  .fetched-card {{ background:{GOLD_LIGHT}; border:1px solid #e8d9b0; border-left:4px solid {GOLD}; border-radius:0 10px 10px 0; padding:1rem 1.25rem; margin:1rem 0 0.5rem; font-size:0.9rem; line-height:1.7; }}
  .fetched-card .fc-label {{ font-weight:700; color:{GOLD_DARK}; font-size:0.75rem; text-transform:uppercase; letter-spacing:0.1em; }}
  .fetched-card .fc-value {{ color:{PRIMARY}; }}
  .upgrade-card {{ background:#fff; border:2px solid {GOLD_MID}; border-top:4px solid {GOLD}; border-radius:14px; padding:2rem; margin:1.5rem 0; text-align:center; }}
  .upgrade-card .uc-eyebrow {{ font-size:0.75rem; font-weight:700; letter-spacing:0.15em; text-transform:uppercase; color:{GOLD_DARK}; margin-bottom:0.4rem; }}
  .upgrade-card h3 {{ font-size:1.4rem; font-weight:900; color:{PRIMARY}; margin:0 0 0.3rem; }}
  .upgrade-card .uc-sub {{ color:#666; font-size:0.92rem; margin-bottom:1.2rem; }}
  .upgrade-card ul {{ text-align:left; display:inline-block; list-style:none; padding:0; margin:0 auto 1.5rem; }}
  .upgrade-card ul li {{ padding:3px 0; font-size:0.92rem; color:#333; }}
  .upgrade-card ul li::before {{ content:"✓ "; color:{GOLD_DARK}; font-weight:700; }}
  .price-badge {{ display:inline-block; background:{PRIMARY}; color:{GOLD}; font-size:1.6rem; font-weight:900; padding:6px 28px; border-radius:40px; margin-bottom:1.2rem; border:2px solid {GOLD}; }}
  .divider-or {{ display:flex; align-items:center; gap:12px; color:#bbb; font-size:0.85rem; margin:1.2rem 0; }}
  .divider-or::before, .divider-or::after {{ content:""; flex:1; border-top:1px solid #e0e0e0; }}
  .coupon-section {{ background:{GOLD_BG}; border:1px solid #e8d78a; border-left:4px solid {GOLD}; border-radius:0 8px 8px 0; padding:1rem 1.25rem; margin-bottom:1rem; }}
  .coupon-section .label {{ font-weight:700; font-size:0.9rem; color:#5C4A00; margin-bottom:0.4rem; }}
  .success-banner {{ background:{GREEN_BG}; border:1px solid #b8ddc8; border-left:4px solid {GREEN}; border-radius:0 8px 8px 0; padding:1rem 1.25rem; margin-bottom:1.5rem; font-size:0.92rem; color:#155724; }}
  .bookmark-tip {{ background:#EEF4FF; border:1px solid #c5d8f8; border-left:4px solid #3B6FD4; border-radius:0 8px 8px 0; padding:0.85rem 1.1rem; margin-bottom:1.5rem; font-size:0.88rem; color:#1a3a6b; }}
  .stButton > button {{ background:{PRIMARY} !important; color:{GOLD} !important; font-weight:700 !important; padding:0.65rem 2rem !important; border-radius:8px !important; border:2px solid {GOLD} !important; font-size:1rem !important; width:100%; }}
  .stButton > button:hover {{ background:{GOLD} !important; color:{PRIMARY} !important; }}
  .stLinkButton > a {{ background:{PRIMARY} !important; color:{GOLD} !important; font-weight:700 !important; border-radius:8px !important; font-size:1rem !important; border:2px solid {GOLD} !important; }}
  .footer-note {{ text-align:center; color:#aaa; font-size:0.78rem; margin-top:3rem; padding-top:1rem; border-top:1px solid {GOLD_MID}; }}
</style>
""", unsafe_allow_html=True)

# ── AI Prompt ─────────────────────────────────────────────────────────────────
AMAZON_SYSTEM_PROMPT = """You are the Amazon Product Page Assessor for The Writing Wives Genre & Trope AI Classroom.

Assess the complete Amazon product page for a book. You will receive: title, subtitle, genre, blurb, star rating, review count, and price. Analyse each element's effectiveness for Amazon discovery and conversion.

Return ONLY valid JSON. No markdown, no code fences. Exact structure:

{
  "summary": {
    "title_subtitle": "✅ Strong",
    "blurb_effectiveness": "⚠️ Needs Work",
    "review_health": "✅ Strong",
    "price_positioning": "⚠️ Needs Work",
    "discoverability": "❌ Missing",
    "conversion_readiness": "⚠️ Needs Work"
  },
  "categories": {
    "title_subtitle":       {"rating": "✅", "observations": ["..."], "suggested_fix": null},
    "blurb_effectiveness":  {"rating": "⚠️", "observations": ["..."], "suggested_fix": "..."},
    "review_health":        {"rating": "✅", "observations": ["..."], "suggested_fix": null},
    "price_positioning":    {"rating": "⚠️", "observations": ["..."], "suggested_fix": "..."},
    "discoverability":      {"rating": "❌", "observations": ["..."], "suggested_fix": "..."},
    "conversion_readiness": {"rating": "⚠️", "observations": ["..."], "suggested_fix": "..."}
  },
  "overall_score": 68,
  "priority_fixes": ["Fix 1 — specific action.", "Fix 2 — specific action.", "Fix 3 — specific action."],
  "star_rating_assessment": {
    "current": 4.1,
    "target": 4.2,
    "status": "Below target",
    "advice": "Specific advice on improving or protecting the rating."
  },
  "keyword_suggestions": ["keyword1", "keyword2", "keyword3", "keyword4", "keyword5"],
  "category_suggestions": ["Fiction > Genre Fiction > Time Travel", "Young Adult > Science Fiction & Fantasy"]
}

Assessment rules:
- title_subtitle: Is the title genre-signalling? Does the subtitle add keywords or promise? Is it Amazon-search-ready?
- blurb_effectiveness: Is the blurb a trope delivery system? Does it hook, raise stakes, and have a GEO end section?
- review_health: Amazon's A9 algorithm favours 4.2+. Flag anything below 4.2 as a critical issue. Also assess review count velocity.
- price_positioning: Is the price appropriate for genre, format, and career stage?
- discoverability: Keywords in title/subtitle/description? Is the book findable through Amazon search?
- conversion_readiness: Would a reader who lands on this page buy? Assess overall page impression.

overall_score: 0–100. 80+ is conversion-ready. 60–79 needs targeted work. Below 60 needs significant overhaul.
Be specific. Every ❌ or ⚠️ MUST have a concrete fix.
keyword_suggestions: 5 specific long-tail keywords the author should add to their KDP keyword fields.
category_suggestions: 2 specific Amazon category paths that would help this book rank."""

def call_amazon_assessment(title, subtitle, genre, blurb, star_rating, review_count, price, categories_current):
    api_key = get_secret("OPENROUTER_API_KEY")
    if not api_key:
        st.error("API key not configured.")
        st.stop()
    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
    model  = get_secret("MODEL", "anthropic/claude-sonnet-4-5")
    parts  = []
    if title:             parts.append(f"Title: {title}")
    if subtitle:          parts.append(f"Subtitle: {subtitle}")
    if genre:             parts.append(f"Genre: {genre}")
    if star_rating:       parts.append(f"Current star rating: {star_rating}")
    if review_count:      parts.append(f"Number of reviews: {review_count}")
    if price:             parts.append(f"Current price: {price}")
    if categories_current: parts.append(f"Current Amazon categories: {categories_current}")
    if blurb:             parts.append(f"\nBlurb / Product Description:\n\n{blurb}")
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": AMAZON_SYSTEM_PROMPT},
            {"role": "user",   "content": "\n".join(parts)},
        ],
        temperature=0.4,
        max_tokens=3000,
    )
    raw = response.choices[0].message.content.strip()
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    return json.loads(raw.strip())

def rating_badge(r):
    if "✅" in r:   return '<span class="r-strong">✅ Strong</span>'
    elif "⚠️" in r: return '<span class="r-warning">⚠️ Needs Work</span>'
    else:            return '<span class="r-missing">❌ Missing</span>'

def render_full_assessment(data, book_title):
    score = data.get("overall_score", 0)
    if score >= 80:
        status_html = f'<div class="score-status" style="color:{GOLD};">✅ Conversion-Ready</div>'
    elif score >= 60:
        status_html = f'<div class="score-status" style="color:#F6A623;">⚠️ Functional — Needs Work</div>'
    else:
        status_html = f'<div class="score-status" style="color:#E74C3C;">❌ Needs Significant Overhaul</div>'

    st.markdown(f"""
    <div class="score-big">
      <div class="score-num">{score}</div>
      <div class="score-label">Amazon Page Score out of 100</div>
      {status_html}
    </div>
    """, unsafe_allow_html=True)

    star_data = data.get("star_rating_assessment", {})
    if star_data:
        current = star_data.get("current", 0)
        status  = star_data.get("status", "")
        advice  = star_data.get("advice", "")
        if "Below" in status:
            box_class, icon = "star-box-bad", "🚨"
        elif "At target" in status:
            box_class, icon = "star-box-warn", "⚠️"
        else:
            box_class, icon = "star-box-good", "✅"
        st.markdown(f"""
        <div class="{box_class}">
          {icon} <strong>Star Rating: {current} ★</strong> — {status} (target: 4.2+)<br>
          <span style="font-size:0.88rem;">{advice}</span>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div class="section-h2">Assessment Summary</div>', unsafe_allow_html=True)
    summary = data.get("summary", {})
    grid = '<div class="sum-grid">'
    for key, label in CATEGORY_LABELS.items():
        grid += f'<div class="sum-cell"><span class="sum-label">{label}</span>{rating_badge(summary.get(key,""))}</div>'
    grid += '</div>'
    st.markdown(grid, unsafe_allow_html=True)

    st.markdown('<div class="section-h2">Detailed Findings</div>', unsafe_allow_html=True)
    cats = data.get("categories", {})
    for key, label in CATEGORY_LABELS.items():
        cat = cats.get(key, {})
        if not cat: continue
        st.markdown(f'<div class="cat-heading">{label} &nbsp; {rating_badge(cat.get("rating",""))}</div>', unsafe_allow_html=True)
        for obs in cat.get("observations", []):
            st.markdown(f"- {obs}")
        if cat.get("suggested_fix"):
            st.markdown(f'<div class="fix-box"><strong>Suggested fix:</strong> {cat["suggested_fix"]}</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-h2">Top 3 Priority Fixes</div>', unsafe_allow_html=True)
    for fix in data.get("priority_fixes", []):
        st.markdown(f'<div class="priority-box">{fix}</div>', unsafe_allow_html=True)

    keywords = data.get("keyword_suggestions", [])
    if keywords:
        st.markdown('<div class="section-h2">🔍 Keyword Suggestions for KDP</div>', unsafe_allow_html=True)
        st.caption("Add these to your KDP keyword fields to improve Amazon search ranking.")
        kw_html = "".join(f'<span class="keyword-tag">{k}</span>' for k in keywords)
        st.markdown(f'<div style="margin:0.5rem 0;">{kw_html}</div>', unsafe_allow_html=True)

    cat_suggestions = data.get("category_suggestions", [])
    if cat_suggestions:
        st.markdown('<div class="section-h2">📂 Category Suggestions</div>', unsafe_allow_html=True)
        st.caption("Request these category placements in your KDP dashboard.")
        for c in cat_suggestions:
            st.markdown(f'<div class="category-item">📂 {c}</div>', unsafe_allow_html=True)

    st.divider()
    st.markdown("**Download your assessment:**")
    ca, cb = st.columns(2)
    with ca:
        st.download_button("⬇ Download as JSON",
            data=json.dumps(data, indent=2, ensure_ascii=False).encode(),
            file_name=f"amazon-assessment-{(book_title or 'report').replace(' ','-').lower()}.json",
            mime="application/json", key="dl_amazon_json")
    with cb:
        if DOCX_AVAILABLE:
            st.download_button("⬇ Download as Word Doc",
                data=generate_amazon_docx(data, book_title),
                file_name=f"Amazon Assessment - {book_title or 'Report'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_amazon_docx")

def render_teaser(data):
    score = data.get("overall_score", 0)
    summary = data.get("summary", {})
    issue_count = sum(1 for v in summary.values() if "⚠️" in v or "❌" in v)
    star_data = data.get("star_rating_assessment", {})
    star_warning = ""
    if star_data and "Below" in star_data.get("status",""):
        star_warning = f" Your star rating ({star_data.get('current','')}) is below the 4.2 target — this is flagged in the full report."
    st.markdown(f"""
    <div style="background:{GOLD_LIGHT};border:2px solid {GOLD_MID};border-left:5px solid {GOLD};border-radius:0 10px 10px 0;padding:1.1rem 1.3rem;margin:1rem 0;">
      <strong>Your Amazon page scored {score}/100</strong> — {issue_count} area{'s' if issue_count != 1 else ''} need attention.{star_warning}
      Unlock the full report for specific fixes.
    </div>
    """, unsafe_allow_html=True)
    st.markdown('<div class="section-h2">Assessment Summary</div>', unsafe_allow_html=True)
    grid = '<div class="sum-grid">'
    for key, label in CATEGORY_LABELS.items():
        grid += f'<div class="sum-cell"><span class="sum-label">{label}</span>{rating_badge(summary.get(key,""))}</div>'
    grid += '</div>'
    st.markdown(grid, unsafe_allow_html=True)

def generate_amazon_docx(data, book_title):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    h = doc.add_heading('Amazon Product Page Assessment', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if book_title:
        p = doc.add_paragraph(book_title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(13)
        p.runs[0].font.bold = True
    doc.add_paragraph(f"Overall Score: {data.get('overall_score', 0)}/100")
    star = data.get("star_rating_assessment", {})
    if star:
        doc.add_paragraph(f"Star Rating: {star.get('current','')} ★ — {star.get('status','')} (target: 4.2+)")
        doc.add_paragraph(f"Advice: {star.get('advice','')}")
    doc.add_heading("Assessment Summary", 1)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = "Category"
    t.rows[0].cells[1].text = "Rating"
    for key, label in CATEGORY_LABELS.items():
        r = t.add_row().cells
        r[0].text = label
        r[1].text = data.get("summary", {}).get(key, "")
    doc.add_heading("Detailed Findings", 1)
    for key, label in CATEGORY_LABELS.items():
        cat = data.get("categories", {}).get(key, {})
        if not cat: continue
        doc.add_heading(f"{label}  {cat.get('rating','')}", 2)
        for obs in cat.get("observations", []):
            doc.add_paragraph(obs, style='List Bullet')
        if cat.get("suggested_fix"):
            p = doc.add_paragraph(f"Fix: {cat['suggested_fix']}")
            p.runs[0].font.italic = True
    doc.add_heading("Top 3 Priority Fixes", 1)
    for fix in data.get("priority_fixes", []):
        doc.add_paragraph(fix, style='List Number')
    if data.get("keyword_suggestions"):
        doc.add_heading("Keyword Suggestions for KDP", 1)
        for kw in data["keyword_suggestions"]:
            doc.add_paragraph(kw, style='List Bullet')
    if data.get("category_suggestions"):
        doc.add_heading("Category Suggestions", 1)
        for c in data["category_suggestions"]:
            doc.add_paragraph(c, style='List Bullet')
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def show_upgrade_card():
    price = get_secret("AMAZON_PRICE_DISPLAY", "$19")
    sub   = get_secret("AMAZON_PRICE_SUBTITLE", "per report · full page review")
    link  = affiliate_payment_link(get_secret("STRIPE_AMAZON_LINK", ""))
    st.markdown(f"""
    <div class="upgrade-card">
      <div class="uc-eyebrow">Unlock Your Full Assessment</div>
      <h3>Every element of your Amazon page — scored and fixed</h3>
      <p class="uc-sub">You've seen your scores. Now get the specific fixes that will improve ranking and conversion.</p>
      <ul>
        <li>6-category detailed findings with fixes</li>
        <li>Star rating health check and advice</li>
        <li>Top 3 priority changes (ranked by impact)</li>
        <li>5 keyword suggestions for your KDP fields</li>
        <li>2 specific Amazon category recommendations</li>
        <li>Downloadable Word doc</li>
      </ul>
      <div class="price-badge">{price}</div>
      <div style="color:#888;font-size:0.85rem;margin-bottom:1.2rem;">{sub}</div>
    </div>
    """, unsafe_allow_html=True)
    if link:
        st.link_button(f"Unlock Full Assessment — {price} →", link, use_container_width=True)

# ── URL verification ──────────────────────────────────────────────────────────
if not is_authenticated():
    sid = st.query_params.get("amazon_session_id")
    if sid and not st.session_state.get("amazon_stripe_checked"):
        st.session_state["amazon_stripe_checked"] = True
        with st.spinner("Verifying your payment..."):
            ok, email = verify_stripe_session(sid)
        if ok:
            grant_access(reason="stripe", email=email)
            st.rerun()
        else:
            st.warning("Payment couldn't be verified yet — try refreshing in a moment.")

# ── Header ────────────────────────────────────────────────────────────────────
if st.button("← Back to All Tools", key="back_home"):
    st.switch_page("home.py")

_logo_b64 = get_logo_b64()
_logo_html = f'<div class="logo-wrap"><img src="data:image/png;base64,{_logo_b64}" alt="The Writing Wives"></div>' if _logo_b64 else ""
st.markdown(f"""
<div class="header-block">
  {_logo_html}
  <div class="eyebrow">Genre &amp; Trope AI Classroom</div>
  <h1>Amazon Page Assessment</h1>
  <p class="sub">Paste your Amazon link — we'll pull the page automatically and score every element.</p>
</div>
""", unsafe_allow_html=True)

# ── Access banners ────────────────────────────────────────────────────────────
if is_authenticated():
    reason = st.session_state.get("amazon_access_reason", "")
    email  = st.session_state.get("amazon_access_email", "")
    if reason == "stripe":
        st.markdown(f'<div class="success-banner">✅ <strong>Payment confirmed{f" for {email}" if email else ""}.</strong> Full access unlocked.</div>', unsafe_allow_html=True)
        st.markdown('<div class="bookmark-tip">🔖 <strong>Bookmark this page</strong> — your URL contains your access token.</div>', unsafe_allow_html=True)
    elif reason in ("coupon", "lifetime"):
        st.markdown('<div class="success-banner">🎓 <strong>Access granted.</strong> Run as many assessments as you need.</div>', unsafe_allow_html=True)

# ── Coupon entry (always rendered so error persists on rerun) ─────────────────
if not is_authenticated():
    st.markdown('<div class="coupon-section"><div class="label">🎓 Writing Wives Skool Member? Enter your coupon for free access.</div></div>', unsafe_allow_html=True)
    if "amazon_coupon_val" not in st.session_state:
        st.session_state["amazon_coupon_val"] = ""
    coupon_col, btn_col = st.columns([3, 1])
    with coupon_col:
        st.text_input("Coupon code", placeholder="e.g. WRITINGWIVES",
                      label_visibility="collapsed", key="amazon_coupon_val")
    with btn_col:
        apply_coupon = st.button("Apply →", key="amazon_coupon_btn", use_container_width=True)
    if apply_coupon:
        entered = st.session_state.get("amazon_coupon_val", "").strip()
        if entered and check_coupon(entered):
            grant_access(reason="coupon")
            st.rerun()
        else:
            st.error("That coupon code wasn't found. Double-check the spelling and try again.")

# ── Main form ─────────────────────────────────────────────────────────────────
with st.form("amazon_url_form"):
    amazon_url = st.text_input(
        "Amazon Book URL",
        placeholder="https://www.amazon.com/dp/B0753LBFQ7",
        help="Paste any Amazon book page URL — we'll extract the ASIN and pull the page automatically.",
    )
    genre = st.text_input(
        "Genre / Sub-Genre",
        placeholder="e.g. Small-Town Romance, Military Sci-Fi, Cozy Mystery",
        help="Tell us the genre so we can calibrate the assessment correctly.",
    )
    submit_label = "Assess My Amazon Page →" if is_authenticated() else "Get My Free Score →"
    submitted = st.form_submit_button(submit_label)

# ── On submit ─────────────────────────────────────────────────────────────────
if submitted:
    url_input = amazon_url.strip()

    if not url_input:
        st.warning("Please paste your Amazon book URL to continue.")
        st.stop()

    if not extract_asin(url_input):
        st.error(
            "That doesn't look like a valid Amazon book URL. "
            "It should contain `/dp/` followed by the ASIN, "
            "e.g. `https://www.amazon.com/dp/B0753LBFQ7`"
        )
        st.stop()

    # ── Scrape ────────────────────────────────────────────────────────────────
    with st.spinner("Fetching your Amazon page..."):
        scraped = scrape_amazon_page(url_input)

    if not scraped["success"]:
        # Scrape failed — show error and fall back to manual form
        st.warning(f"⚠️ {scraped['error']}")
        st.markdown("**No problem — fill in your details below and we'll run the assessment from there.**")

        with st.form("amazon_manual_fallback"):
            col1, col2 = st.columns(2)
            with col1:
                fb_title      = st.text_input("Book Title")
                fb_subtitle   = st.text_input("Subtitle (if any)")
                fb_star       = st.number_input("Star Rating", 0.0, 5.0, 0.0, 0.1, format="%.1f")
            with col2:
                fb_price      = st.text_input("Current Price", placeholder="e.g. $4.99")
                fb_reviews    = st.number_input("Number of Reviews", 0, step=1)
                fb_categories = st.text_input("Current Categories (optional)")
            fb_blurb = st.text_area("Your Blurb / Product Description", height=180)
            fb_submit = st.form_submit_button("Run Assessment →")

        if fb_submit:
            if not fb_title.strip() and not fb_blurb.strip():
                st.warning("Please enter at least your book title and blurb.")
            else:
                with st.spinner("Assessing your Amazon page..."):
                    try:
                        result = call_amazon_assessment(
                            fb_title, fb_subtitle, genre, fb_blurb,
                            fb_star if fb_star > 0 else None,
                            fb_reviews if fb_reviews > 0 else None,
                            fb_price, fb_categories,
                        )
                        st.session_state["last_amazon"] = {"result": result, "book_title": fb_title}
                    except json.JSONDecodeError:
                        st.error("The AI returned an unexpected format. Please try again.")
                        st.stop()
                    except Exception as e:
                        st.error("Something went wrong. Please try again.")
                        st.exception(e)
                        st.stop()

                st.success("Assessment complete!")
                st.divider()
                if is_authenticated():
                    render_full_assessment(result, fb_title)
                else:
                    render_teaser(result)
                    show_upgrade_card()

    else:
        # ── Show what was fetched ──────────────────────────────────────────────
        rows = []
        if scraped["title"]:   rows.append(("Title",         scraped["title"]))
        if scraped["author"]:  rows.append(("Author",        scraped["author"]))
        if scraped["price"]:   rows.append(("Price",         scraped["price"]))
        if scraped["star_rating"] is not None:
            rows.append(("Star Rating",  f"{scraped['star_rating']} ★"))
        if scraped["review_count"] is not None:
            rows.append(("Reviews",      f"{scraped['review_count']:,}"))
        if scraped["categories"]:
            rows.append(("Categories",   scraped["categories"]))

        if rows:
            row_html = "".join(
                f'<div><span class="fc-label">{label}</span><br>'
                f'<span class="fc-value">{value}</span></div>'
                for label, value in rows
            )
            st.markdown(f'<div class="fetched-card">✅ <strong>Page fetched successfully</strong><br><br>{row_html}</div>', unsafe_allow_html=True)

        # ── Run assessment ────────────────────────────────────────────────────
        with st.spinner("Assessing your Amazon page..."):
            try:
                result = call_amazon_assessment(
                    scraped["title"],
                    scraped["subtitle"],
                    genre,
                    scraped["blurb"],
                    scraped["star_rating"],
                    scraped["review_count"],
                    scraped["price"],
                    scraped["categories"],
                )
                st.session_state["last_amazon"] = {
                    "result": result,
                    "book_title": scraped["title"],
                }
            except json.JSONDecodeError:
                st.error("The AI returned an unexpected format. Please try again.")
                st.stop()
            except Exception as e:
                st.error("Something went wrong. Please try again.")
                st.exception(e)
                st.stop()

        st.success("Assessment complete!")
        st.divider()

        if is_authenticated():
            render_full_assessment(result, scraped["title"])
        else:
            render_teaser(result)
            show_upgrade_card()

# ── Coupon unlock with cached result ─────────────────────────────────────────
elif is_authenticated() and st.session_state.get("last_amazon") and st.session_state.get("amazon_access_reason") == "coupon":
    prior = st.session_state["last_amazon"]
    st.success("✅ Access granted! Here's your full assessment:")
    st.divider()
    render_full_assessment(prior["result"], prior["book_title"])
    del st.session_state["last_amazon"]

# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="footer-note">
  The Writing Wives · Genre &amp; Trope AI Classroom ·
  <a href="https://thewritingwives.com" style="color:#999;">thewritingwives.com</a>
</div>
""", unsafe_allow_html=True)
