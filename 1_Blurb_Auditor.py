import streamlit as st
from openai import OpenAI
import json
import re
import io
import base64
from pathlib import Path

try:
    import stripe as stripe_lib
    STRIPE_AVAILABLE = True
except ImportError:
    STRIPE_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── Page config ───────────────────────────────────────────────────────────────


# ── Colours — Writing Wives brand (black + gold) ──────────────────────────────
PRIMARY      = "#1A1A1A"   # near-black
PRIMARY_DARK = "#000000"
PRIMARY_MID  = "#2E2E2E"
GOLD         = "#D4B36E"   # exact brand gold
GOLD_DARK    = "#A8863A"
GOLD_LIGHT   = "#FAF4E6"
GOLD_MID     = "#F5EDD6"
GOLD_BG      = "#FFFBEC"
GREEN        = "#1A7A4A"
GREEN_BG     = "#D4EDDA"

# Keep these aliases so existing code that refs PLUM still works
PLUM       = PRIMARY
PLUM_DARK  = PRIMARY_DARK
PLUM_LIGHT = GOLD_MID
PLUM_MID   = GOLD_LIGHT

# ── Logo loader ───────────────────────────────────────────────────────────────
def get_logo_b64():
    logo_path = Path(__file__).parent / "logo.png"
    if logo_path.exists():
        with open(logo_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    return None

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
  html, body, [class*="css"] {{ font-family: 'Segoe UI', Arial, sans-serif; }}
  .block-container {{ max-width: 760px; padding-top: 1.5rem; }}

  /* Header */
  .header-block {{
    background: {PRIMARY};
    padding: 2rem 2.5rem 1.6rem;
    border-radius: 12px;
    margin-bottom: 2rem;
    border-bottom: 3px solid {GOLD};
  }}
  .header-block .logo-wrap {{ text-align:center; margin-bottom:1.1rem; }}
  .header-block .logo-wrap img {{ height:160px; max-width:100%; object-fit:contain; }}
  .header-block .eyebrow {{ color:{GOLD}; font-size:0.72rem; font-weight:700; letter-spacing:0.18em; text-transform:uppercase; margin:0 0 0.3rem; text-align:center; }}
  .header-block h1 {{ color:#fff; font-size:2.1rem; font-weight:900; margin:0 0 0.35rem; line-height:1.15; text-align:center; }}
  .header-block .sub {{ color:#aaa; font-size:0.95rem; margin:0; text-align:center; }}

  /* Rating badges */
  .r-strong  {{ background:#D4EDDA; color:#155724; padding:3px 12px; border-radius:20px; font-weight:700; font-size:0.88rem; display:inline-block; }}
  .r-warning {{ background:#FFF3CD; color:#856404; padding:3px 12px; border-radius:20px; font-weight:700; font-size:0.88rem; display:inline-block; }}
  .r-missing {{ background:#F8D7DA; color:#721C24; padding:3px 12px; border-radius:20px; font-weight:700; font-size:0.88rem; display:inline-block; }}

  /* Audit display */
  .cat-heading {{ color:{PRIMARY}; font-size:1.1rem; font-weight:700; border-bottom:2px solid {GOLD}; padding-bottom:4px; margin:1.6rem 0 0.6rem; }}
  .section-h2  {{ color:{PRIMARY}; font-size:1.3rem; font-weight:900; border-bottom:3px solid {GOLD}; padding-bottom:6px; margin:2rem 0 1rem; }}
  .fix-box     {{ background:{GOLD_LIGHT}; border-left:4px solid {GOLD}; padding:0.75rem 1rem; border-radius:0 8px 8px 0; margin-top:0.5rem; font-size:0.92rem; line-height:1.65; }}
  .rewrite-box {{ background:{GOLD_LIGHT}; border-left:5px solid {GOLD}; padding:1.1rem 1.3rem; border-radius:0 10px 10px 0; margin:0.5rem 0 0.3rem; font-style:italic; line-height:1.8; font-size:0.95rem; white-space:pre-wrap; }}
  .rewrite-notes {{ font-size:0.82rem; color:#666; margin:0 0 1.5rem; font-style:italic; }}
  .priority-box  {{ background:{GOLD_BG}; border-left:5px solid {GOLD_DARK}; padding:0.9rem 1.1rem; border-radius:0 10px 10px 0; margin:0.4rem 0; font-size:0.93rem; line-height:1.6; }}
  .blurb-submitted {{ background:#f9f9f9; border:1px solid #e0d9c8; padding:1rem 1.3rem; border-radius:8px; font-style:italic; line-height:1.75; color:#333; font-size:0.95rem; white-space:pre-wrap; }}
  .sum-grid {{ display:grid; grid-template-columns:1fr 1fr; gap:8px; margin:0.75rem 0 1.5rem; }}
  .sum-cell {{ background:{GOLD_MID}; border-radius:8px; padding:10px 14px; display:flex; align-items:center; justify-content:space-between; font-size:0.9rem; border:1px solid #E8D9B0; }}
  .sum-cell .sum-label {{ font-weight:600; color:#333; }}
  .trope-table {{ width:100%; border-collapse:collapse; font-size:0.9rem; margin:0.5rem 0 1rem; }}
  .trope-table th {{ background:{PRIMARY}; color:{GOLD}; padding:7px 10px; text-align:left; }}
  .trope-table td {{ padding:7px 10px; border-bottom:1px solid #eee; vertical-align:top; }}
  .trope-table tr:nth-child(even) td {{ background:{GOLD_LIGHT}; }}
  .t-present {{ color:#155724; font-weight:700; }}
  .t-implied {{ color:#856404; font-weight:700; }}
  .t-missing {{ color:#721C24; font-weight:700; }}

  /* ── Upgrade wall ── */
  .upgrade-wall {{
    position: relative;
    margin: 2rem 0 0;
  }}
  .upgrade-blur {{
    filter: blur(5px);
    pointer-events: none;
    user-select: none;
    opacity: 0.55;
    max-height: 260px;
    overflow: hidden;
  }}
  .upgrade-card {{
    background: #fff;
    border: 2px solid {GOLD_MID};
    border-top: 4px solid {GOLD};
    border-radius: 14px;
    padding: 2rem 2rem 1.5rem;
    margin: 1.5rem 0;
    text-align: center;
  }}
  .upgrade-card .uc-eyebrow {{
    font-size: 0.75rem;
    font-weight: 700;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    color: {GOLD_DARK};
    margin-bottom: 0.4rem;
  }}
  .upgrade-card h3 {{
    font-size: 1.4rem;
    font-weight: 900;
    color: {PRIMARY};
    margin: 0 0 0.3rem;
  }}
  .upgrade-card .uc-sub {{
    color: #666;
    font-size: 0.92rem;
    margin-bottom: 1.2rem;
  }}
  .upgrade-card ul {{
    text-align: left;
    display: inline-block;
    list-style: none;
    padding: 0;
    margin: 0 auto 1.5rem;
  }}
  .upgrade-card ul li {{
    padding: 3px 0;
    font-size: 0.92rem;
    color: #333;
  }}
  .upgrade-card ul li::before {{ content: "✓ "; color:{GOLD_DARK}; font-weight:700; }}
  .price-badge {{
    display: inline-block;
    background: {PRIMARY};
    color: {GOLD};
    font-size: 1.6rem;
    font-weight: 900;
    padding: 6px 28px;
    border-radius: 40px;
    margin-bottom: 1.2rem;
    border: 2px solid {GOLD};
  }}
  .divider-or {{
    display:flex; align-items:center; gap:12px;
    color:#bbb; font-size:0.85rem; margin:1.2rem 0;
  }}
  .divider-or::before, .divider-or::after {{ content:""; flex:1; border-top:1px solid #e0e0e0; }}
  .coupon-section {{
    background:{GOLD_BG}; border:1px solid #e8d78a;
    border-left:4px solid {GOLD}; border-radius:0 8px 8px 0;
    padding:1rem 1.25rem; margin-bottom:1rem;
  }}
  .coupon-section .label {{ font-weight:700; font-size:0.9rem; color:#5C4A00; margin-bottom:0.4rem; }}

  /* Banners */
  .success-banner {{
    background:{GREEN_BG}; border:1px solid #b8ddc8;
    border-left:4px solid {GREEN}; border-radius:0 8px 8px 0;
    padding:1rem 1.25rem; margin-bottom:1.5rem;
    font-size:0.92rem; color:#155724;
  }}
  .bookmark-tip {{
    background:#EEF4FF; border:1px solid #c5d8f8;
    border-left:4px solid #3B6FD4; border-radius:0 8px 8px 0;
    padding:0.85rem 1.1rem; margin-bottom:1.5rem;
    font-size:0.88rem; color:#1a3a6b;
  }}
  .issues-banner {{
    background:{GOLD_LIGHT}; border:2px solid {GOLD_MID};
    border-left:5px solid {GOLD}; border-radius:0 10px 10px 0;
    padding:1.1rem 1.3rem; margin:1rem 0;
    font-size:0.95rem; color:{PRIMARY};
  }}
  .issues-banner strong {{ font-size:1.05rem; }}

  /* Button */
  .stButton > button {{
    background:{PRIMARY} !important; color:{GOLD} !important;
    font-weight:700 !important; padding:0.65rem 2rem !important;
    border-radius:8px !important; border:2px solid {GOLD} !important;
    font-size:1rem !important; width:100%;
  }}
  .stButton > button:hover {{ background:{GOLD} !important; color:{PRIMARY} !important; }}
  .stLinkButton > a {{
    background:{PRIMARY} !important; color:{GOLD} !important;
    font-weight:700 !important; border-radius:8px !important;
    font-size:1rem !important; border:2px solid {GOLD} !important;
  }}

  .footer-note {{ text-align:center; color:#aaa; font-size:0.78rem; margin-top:3rem; padding-top:1rem; border-top:1px solid {GOLD_MID}; }}

  /* ── FB Ad Package upsell card ── */
  .ad-upgrade-card {{
    background: {PRIMARY};
    border: 2px solid {GOLD};
    border-radius: 14px;
    padding: 2rem 2rem 1.5rem;
    margin: 2rem 0;
    text-align: center;
  }}
  .ad-upgrade-card .uc-eyebrow {{
    font-size: 0.75rem; font-weight: 700; letter-spacing: 0.15em;
    text-transform: uppercase; color: {GOLD}; margin-bottom: 0.4rem;
  }}
  .ad-upgrade-card h3 {{ font-size: 1.4rem; font-weight: 900; color: #fff; margin: 0 0 0.3rem; }}
  .ad-upgrade-card .uc-sub {{ color: #bbb; font-size: 0.92rem; margin-bottom: 1.2rem; }}
  .ad-upgrade-card ul {{
    text-align: left; display: inline-block; list-style: none;
    padding: 0; margin: 0 auto 1.5rem;
  }}
  .ad-upgrade-card ul li {{ padding: 3px 0; font-size: 0.92rem; color: #ddd; }}
  .ad-upgrade-card ul li::before {{ content: "✓ "; color:{GOLD}; font-weight:700; }}
  .ad-price-badge {{
    display: inline-block; background: {GOLD}; color: {PRIMARY};
    font-size: 1.6rem; font-weight: 900; padding: 6px 28px;
    border-radius: 40px; margin-bottom: 1.2rem;
  }}

  /* ── FB Ad Package results ── */
  .ad-section-header {{
    background: {PRIMARY}; border-radius: 12px;
    padding: 1.2rem 2rem; margin: 2.5rem 0 1.5rem;
    border-bottom: 3px solid {GOLD}; text-align: center;
  }}
  .ad-section-header .ad-eyebrow {{
    color: {GOLD}; font-size: 0.72rem; font-weight: 700;
    letter-spacing: 0.18em; text-transform: uppercase; margin: 0 0 0.25rem;
  }}
  .ad-section-header h2 {{ color: #fff; font-size: 1.5rem; font-weight: 900; margin: 0; }}
  .best-pick-badge {{
    background: {GOLD}; color: {PRIMARY}; font-weight: 700; font-size: 0.75rem;
    padding: 2px 10px; border-radius: 20px; display: inline-block;
    margin-bottom: 0.35rem; letter-spacing: 0.05em; text-transform: uppercase;
  }}
  .headline-card {{
    background: {GOLD_LIGHT}; border: 1px solid #E8D9B0;
    border-left: 4px solid {GOLD}; border-radius: 0 8px 8px 0;
    padding: 0.85rem 1.1rem; margin: 0.5rem 0;
  }}
  .headline-card.best {{ border-left: 4px solid {GOLD_DARK}; background: {GOLD_MID}; }}
  .headline-card .hl-text {{ font-weight: 700; font-size: 1.05rem; color: {PRIMARY}; margin-bottom: 0.25rem; }}
  .headline-card .hl-meta {{ font-size: 0.8rem; color: #888; }}
  .pt-card {{
    background: #fff; border: 1px solid #e0d9c8;
    border-radius: 8px; padding: 1rem 1.3rem; margin: 0.75rem 0;
  }}
  .pt-card.best {{ border: 2px solid {GOLD}; }}
  .pt-label {{
    font-weight: 700; font-size: 0.82rem; color: {GOLD_DARK};
    text-transform: uppercase; letter-spacing: 0.1em; margin-bottom: 0.5rem;
  }}
  .pt-text {{ font-size: 0.93rem; line-height: 1.75; color: #333; white-space: pre-wrap; }}
  .pt-notes {{ font-size: 0.8rem; color: #888; margin-top: 0.5rem; font-style: italic; }}
  .desc-card {{
    background: {GOLD_BG}; border-left: 4px solid {GOLD};
    border-radius: 0 8px 8px 0; padding: 0.75rem 1rem; margin: 0.5rem 0;
  }}
  .desc-card.best {{ border-left: 4px solid {GOLD_DARK}; }}
  .desc-text {{ font-size: 0.95rem; color: {PRIMARY}; font-weight: 500; margin-bottom: 0.2rem; }}
  .desc-notes {{ font-size: 0.8rem; color: #888; font-style: italic; }}
  .targeting-box {{
    background: #EEF4FF; border: 1px solid #c5d8f8;
    border-left: 4px solid #3B6FD4; border-radius: 0 8px 8px 0;
    padding: 0.85rem 1.1rem; margin: 0.5rem 0;
    font-size: 0.92rem; color: #1a3a6b; line-height: 1.6;
  }}
</style>
""", unsafe_allow_html=True)

# ── Category config ───────────────────────────────────────────────────────────
CATEGORY_LABELS = {
    "opening_hook":  "1. Opening Hook",
    "structure":     "2. Structure",
    "trope_signals": "3. Trope Signals",
    "tone_style":    "4. Tone & Style",
    "stakes":        "5. Stakes",
    "geo_section":   "6. GEO End Section",
}
# Only the first category is shown in the free teaser
FREE_CATEGORIES = ["opening_hook"]

# ── AI prompt ─────────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are the Blurb Auditor for The Writing Wives Genre & Trope AI Classroom.

A book blurb is a TROPE DELIVERY SYSTEM, not a plot summary. Audit how well a blurb signals genre, surfaces tropes, establishes stakes, hooks readers in the first sentence, and includes GEO (Generative Engine Optimisation) elements for AI discovery tools.

Audit six categories, each rated: ✅ Strong / ⚠️ Needs Work / ❌ Missing

1. Opening Hook — Is the first sentence the most gripping line? Does it signal genre? Does it avoid backstory?
2. Structure — Are the four framework elements present (status quo, inciting incident, character's plan, stakes)?
3. Trope Signals — Are key tropes visible? Any missing or reader-irritating tropes being led with?
4. Tone & Style — Does tone match genre? Is person/tense consistent? Any filler phrases?
5. Stakes — Are consequences specific and named, or vague?
6. GEO End Section — Is there a trope declaration + comp title section for AI recommendation tools?

CRITICAL: Return ONLY valid JSON. No markdown, no code fences, no preamble. Exact structure:

{
  "summary": {
    "opening_hook": "✅ Strong",
    "structure": "⚠️ Needs Work",
    "trope_signals": "✅ Strong",
    "tone_style": "⚠️ Needs Work",
    "stakes": "❌ Missing",
    "geo_section": "❌ Missing"
  },
  "categories": {
    "opening_hook":  { "rating": "✅", "observations": ["Specific obs 1.", "Specific obs 2."], "suggested_fix": null },
    "structure":     { "rating": "⚠️", "observations": ["Specific obs."], "suggested_fix": "Concrete fix." },
    "trope_signals": { "rating": "✅", "observations": ["Specific obs."], "suggested_fix": null },
    "tone_style":    { "rating": "⚠️", "observations": ["Specific obs."], "suggested_fix": "Concrete fix." },
    "stakes":        { "rating": "❌", "observations": ["Specific obs 1.", "Specific obs 2."], "suggested_fix": "Concrete fix." },
    "geo_section":   { "rating": "❌", "observations": ["No GEO section detected."], "suggested_fix": "Add: 'If you love [trope] and [trope], this is your next read. Fans of [COMP TITLE] will devour it.'" }
  },
  "trope_check": null,
  "priority_fixes": ["Fix 1 — specific action.", "Fix 2 — specific action.", "Fix 3 — specific action."],
  "rewrite_a": {
    "label": "Option A — Conservative",
    "text": "Full rewritten blurb preserving author voice. Use \\n\\n between paragraphs.",
    "notes": "1–2 sentence explanation."
  },
  "rewrite_b": {
    "label": "Option B — Bold",
    "text": "Full rewritten blurb with stronger choices. Use \\n\\n between paragraphs.",
    "notes": "1–2 sentence explanation."
  }
}

If tropes provided: "trope_check": [{"trope": "Name", "status": "Present|Implied|Missing", "evidence": "quote or suggestion"}]

Be specific and direct. Every ❌ or ⚠️ MUST have a concrete fix. Match the author's person and tense in both rewrites. Both rewrites must include a strong GEO end section."""


# ── FB Ad Package AI prompt ───────────────────────────────────────────────────
AD_SYSTEM_PROMPT = """You are the Facebook Ad Copy Generator for The Writing Wives Genre & Trope AI Classroom.

Generate complete, ready-to-paste Facebook and Instagram ad copy for a book based on its blurb. You may also receive blurb audit insights — use these to lean into the blurb's strengths and avoid its weaknesses in the ad copy.

Return ONLY valid JSON. No markdown, no code fences, no preamble. Exact structure:

{
  "headlines": [
    {"text": "Headline text here", "chars": 28, "angle": "Brief note on angle: curiosity/trope/stakes/etc"},
    {"text": "Second headline", "chars": 32, "angle": "..."},
    {"text": "Third headline", "chars": 35, "angle": "..."},
    {"text": "Fourth headline", "chars": 29, "angle": "..."},
    {"text": "Fifth headline", "chars": 33, "angle": "..."}
  ],
  "primary_texts": [
    {"label": "Short (under 80 words)", "text": "Full ad copy text...\\n\\nSoft CTA here →", "notes": "Best for cold traffic / mobile scroll"},
    {"label": "Medium (80–150 words)", "text": "Full ad copy text...\\n\\nSoft CTA here →", "notes": "Best for warm audiences / retargeting"},
    {"label": "Long (150–250 words)", "text": "Full ad copy text...\\n\\nSoft CTA here →", "notes": "Best for engaged readers / newsletter lookalikes"}
  ],
  "descriptions": [
    {"text": "Short punchy description under 30 words", "notes": "Angle: trope signal / reader promise / etc"},
    {"text": "Alternative description under 30 words", "notes": "Angle: ..."}
  ],
  "best_picks": {
    "headline": 2,
    "primary_text": 1,
    "description": 0
  },
  "targeting_note": "1–2 sentence suggestion on Facebook audience targeting based on the genre and tropes detected."
}

Headline rules (5 total):
- Under 40 characters EACH — count them carefully
- Never open with the book title
- Lead with the hook, trope, or emotional tension
- Vary the angles across the 5: curiosity / trope-forward / emotional / stakes-driven / reader-promise

Primary text rules:
- First word must NOT be "I"
- Open with a single punchy hook line on its own
- Use short paragraphs and line breaks for mobile readability
- Include a soft CTA at the end (e.g. "Grab it now →" or "Start reading free →")
- Genre-aware and trope-forward — readers should feel the vibe immediately

Description rules:
- Under 30 words each
- One clear reader promise or trope signal per option
- No filler phrases like "a story of" or "an unforgettable tale"

Best picks: Be decisive. Pick the single most likely to perform in each category."""


# ── Auth helpers ──────────────────────────────────────────────────────────────

def verify_stripe_session(session_id):
    if not STRIPE_AVAILABLE:
        return False, None
    try:
        stripe_lib.api_key = st.secrets["STRIPE_SECRET_KEY"]
        session = stripe_lib.checkout.Session.retrieve(session_id)
        if session.payment_status == "paid":
            email = session.customer_details.email if session.customer_details else None
            return True, email
    except Exception:
        pass
    return False, None


def check_coupon(code):
    raw = st.secrets.get("COUPON_CODES", "")
    valid = {c.strip().upper() for c in raw.split(",") if c.strip()}
    return code.strip().upper() in valid


def is_authenticated():
    return st.session_state.get("access_granted", False) or st.session_state.get("lifetime_access", False)


def grant_access(reason="", email=None):
    st.session_state["access_granted"] = True
    st.session_state["access_reason"]  = reason
    if email:
        st.session_state["access_email"] = email


def is_ad_authenticated():
    return st.session_state.get("ad_access_granted", False) or st.session_state.get("lifetime_access", False)


def grant_ad_access(reason="", email=None):
    st.session_state["ad_access_granted"] = True
    st.session_state["ad_access_reason"]  = reason
    if email:
        st.session_state["ad_access_email"] = email


def check_ad_coupon(code):
    # Use AD_COUPON_CODES if configured separately, otherwise fall back to COUPON_CODES
    raw = st.secrets.get("AD_COUPON_CODES", st.secrets.get("COUPON_CODES", ""))
    valid = {c.strip().upper() for c in raw.split(",") if c.strip()}
    return code.strip().upper() in valid


# ── AI calls ──────────────────────────────────────────────────────────────────

def call_openrouter(blurb, genre, tense, tropes, book_title):
    try:
        api_key = st.secrets["OPENROUTER_API_KEY"]
    except Exception:
        st.error("API key not configured. Add OPENROUTER_API_KEY to your Streamlit secrets.")
        st.stop()

    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
    model  = st.secrets.get("MODEL", "anthropic/claude-sonnet-4-5")

    parts = []
    if book_title: parts.append(f"Book title: {book_title}")
    if genre:      parts.append(f"Genre: {genre}")
    if tense:      parts.append(f"POV / Tense: {tense}")
    if tropes:     parts.append(f"Author's top tropes: {tropes}")
    parts.append(f"\nBlurb to audit:\n\n{blurb}")

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": "\n".join(parts)},
        ],
        temperature=0.4,
        max_tokens=4000,
    )
    raw = response.choices[0].message.content.strip()
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    return json.loads(raw.strip())


def call_openrouter_ads(blurb, genre, book_title, audit_data=None):
    """Generate FB ad copy, optionally informed by the blurb audit results."""
    try:
        api_key = st.secrets["OPENROUTER_API_KEY"]
    except Exception:
        st.error("API key not configured. Add OPENROUTER_API_KEY to your Streamlit secrets.")
        st.stop()

    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
    model  = st.secrets.get("MODEL", "anthropic/claude-sonnet-4-5")

    parts = []
    if book_title: parts.append(f"Book title: {book_title}")
    if genre:      parts.append(f"Genre: {genre}")
    parts.append(f"\nBlurb:\n\n{blurb}")

    # Enrich the prompt with audit insights so ad copy plays to the blurb's strengths
    if audit_data:
        summary   = audit_data.get("summary", {})
        strengths = [CATEGORY_LABELS.get(k, k) for k, v in summary.items() if "✅" in v]
        weaknesses = [CATEGORY_LABELS.get(k, k) for k, v in summary.items() if "❌" in v or "⚠️" in v]
        if strengths:
            parts.append(f"\nBlurb strengths to lean into: {', '.join(strengths)}")
        if weaknesses:
            parts.append(f"Blurb weaknesses to work around: {', '.join(weaknesses)}")
        fixes = audit_data.get("priority_fixes", [])
        if fixes:
            parts.append(f"Key issues to avoid reflecting in ad copy: {'; '.join(fixes[:2])}")

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": AD_SYSTEM_PROMPT},
            {"role": "user",   "content": "\n".join(parts)},
        ],
        temperature=0.5,
        max_tokens=3000,
    )
    raw = response.choices[0].message.content.strip()
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    return json.loads(raw.strip())


# ── Render helpers ────────────────────────────────────────────────────────────

def rating_badge(r):
    if "✅" in r:   return '<span class="r-strong">✅ Strong</span>'
    elif "⚠️" in r: return '<span class="r-warning">⚠️ Needs Work</span>'
    else:            return '<span class="r-missing">❌ Missing</span>'


def render_summary_grid(summary):
    grid = '<div class="sum-grid">'
    for key, label in CATEGORY_LABELS.items():
        grid += f'<div class="sum-cell"><span class="sum-label">{label}</span>{rating_badge(summary.get(key,""))}</div>'
    grid += '</div>'
    st.markdown(grid, unsafe_allow_html=True)


def render_category(key, label, cat):
    st.markdown(f'<div class="cat-heading">{label} &nbsp; {rating_badge(cat.get("rating",""))}</div>', unsafe_allow_html=True)
    for obs in cat.get("observations", []):
        st.markdown(f"- {obs}")
    if cat.get("suggested_fix"):
        st.markdown(f'<div class="fix-box"><strong>Suggested fix:</strong> {cat["suggested_fix"]}</div>', unsafe_allow_html=True)


def render_full_audit(data, blurb, book_title):
    """Render the complete audit — for paid/coupon users."""
    st.markdown('<div class="section-h2">Blurb as Submitted</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="blurb-submitted">{blurb}</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-h2">Audit Summary</div>', unsafe_allow_html=True)
    render_summary_grid(data.get("summary", {}))

    st.markdown('<div class="section-h2">Detailed Findings</div>', unsafe_allow_html=True)
    cats = data.get("categories", {})
    for key, label in CATEGORY_LABELS.items():
        cat = cats.get(key, {})
        if cat:
            render_category(key, label, cat)

    trope_check = data.get("trope_check")
    if trope_check:
        st.markdown('<div class="section-h2">Trope Check</div>', unsafe_allow_html=True)
        rows = ""
        for t in trope_check:
            s   = t.get("status", "")
            cls = "t-present" if s == "Present" else ("t-implied" if s == "Implied" else "t-missing")
            rows += f'<tr><td><strong>{t.get("trope","")}</strong></td><td class="{cls}">{s}</td><td>{t.get("evidence","")}</td></tr>'
        st.markdown(f'<table class="trope-table"><thead><tr><th>Trope</th><th>Status</th><th>Evidence / Suggestion</th></tr></thead><tbody>{rows}</tbody></table>', unsafe_allow_html=True)

    st.markdown('<div class="section-h2">Top 3 Priority Fixes</div>', unsafe_allow_html=True)
    for fix in data.get("priority_fixes", []):
        st.markdown(f'<div class="priority-box">{fix}</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-h2">Rewritten Blurb Options</div>', unsafe_allow_html=True)
    for key in ("rewrite_a", "rewrite_b"):
        rw = data.get(key, {})
        if not rw: continue
        st.markdown(f"**{rw.get('label', key)}**")
        st.markdown(f'<div class="rewrite-box">{rw.get("text","").replace(chr(10),"<br>")}</div>', unsafe_allow_html=True)
        if rw.get("notes"):
            st.markdown(f'<div class="rewrite-notes">✏️ {rw["notes"]}</div>', unsafe_allow_html=True)

    # Downloads
    st.divider()
    st.markdown("**Download your audit:**")
    col_a, col_b = st.columns(2)
    with col_a:
        st.download_button(
            "⬇ Download as JSON",
            data=json.dumps(data, indent=2, ensure_ascii=False).encode("utf-8"),
            file_name=f"blurb-audit-{(book_title or 'report').replace(' ','-').lower()}.json",
            mime="application/json",
        )
    with col_b:
        if DOCX_AVAILABLE:
            st.download_button(
                "⬇ Download as Word Doc",
                data=generate_docx(data, blurb, book_title),
                file_name=f"Blurb Audit - {book_title or 'Report'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


def render_teaser_audit(data, blurb):
    """Render the free teaser — summary grid + first category only, rest locked."""

    # Count issues
    summary = data.get("summary", {})
    issue_count = sum(1 for v in summary.values() if "⚠️" in v or "❌" in v)
    missing_count = sum(1 for v in summary.values() if "❌" in v)

    st.markdown('<div class="section-h2">Blurb as Submitted</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="blurb-submitted">{blurb}</div>', unsafe_allow_html=True)

    # Issues banner
    st.markdown(f"""
    <div class="issues-banner">
      <strong>Your blurb has {issue_count} issue{'s' if issue_count != 1 else ''} flagged</strong>
      {'— including <strong>' + str(missing_count) + ' missing element' + ('s' if missing_count != 1 else '') + '</strong>' if missing_count else ''}.
      Here's a preview of your results.
    </div>
    """, unsafe_allow_html=True)

    # Summary grid — always shown
    st.markdown('<div class="section-h2">Audit Summary</div>', unsafe_allow_html=True)
    render_summary_grid(summary)

    # First category — fully revealed
    st.markdown('<div class="section-h2">Detailed Findings — Preview</div>', unsafe_allow_html=True)
    cats = data.get("categories", {})
    first_key  = FREE_CATEGORIES[0]
    first_label = CATEGORY_LABELS[first_key]
    cat = cats.get(first_key, {})
    if cat:
        render_category(first_key, first_label, cat)

    # ── Blurred preview of remaining categories ───────────────────────────────
    locked_preview_html = ""
    for key, label in CATEGORY_LABELS.items():
        if key in FREE_CATEGORIES:
            continue
        rating = summary.get(key, "")
        badge  = rating_badge(rating)
        locked_preview_html += f'<div style="margin:1rem 0;"><div style="color:{PRIMARY};font-weight:700;font-size:1.05rem;border-bottom:2px solid {GOLD};padding-bottom:3px;margin-bottom:6px;">{label} &nbsp; {badge}</div>'
        locked_preview_html += '<div style="color:#444;font-size:0.9rem;line-height:1.6;">Detailed analysis and suggested fix available in the full report.</div></div>'

    locked_preview_html += f"""
    <div style="margin-top:1.5rem;">
      <div style="color:{PRIMARY};font-size:1.3rem;font-weight:900;border-bottom:3px solid {GOLD};padding-bottom:6px;margin-bottom:0.75rem;">Top 3 Priority Fixes</div>
      <div style="background:{GOLD_BG};border-left:5px solid {GOLD_DARK};padding:0.9rem 1.1rem;border-radius:0 10px 10px 0;margin:0.4rem 0;font-size:0.93rem;">Specific priority fix with action steps...</div>
      <div style="background:{GOLD_BG};border-left:5px solid {GOLD_DARK};padding:0.9rem 1.1rem;border-radius:0 10px 10px 0;margin:0.4rem 0;font-size:0.93rem;">Specific priority fix with action steps...</div>
      <div style="background:{GOLD_BG};border-left:5px solid {GOLD_DARK};padding:0.9rem 1.1rem;border-radius:0 10px 10px 0;margin:0.4rem 0;font-size:0.93rem;">Specific priority fix with action steps...</div>
    </div>
    <div style="margin-top:1.5rem;">
      <div style="color:{PRIMARY};font-size:1.3rem;font-weight:900;border-bottom:3px solid {GOLD};padding-bottom:6px;margin-bottom:0.75rem;">Rewritten Blurb Options</div>
      <div style="background:{GOLD_LIGHT};border-left:5px solid {GOLD};padding:1.1rem 1.3rem;border-radius:0 10px 10px 0;margin:0.5rem 0;font-style:italic;line-height:1.8;">Option A — Conservative rewrite ready to copy to your retailer page...</div>
      <div style="background:{GOLD_LIGHT};border-left:5px solid {GOLD};padding:1.1rem 1.3rem;border-radius:0 10px 10px 0;margin:0.5rem 0;font-style:italic;line-height:1.8;">Option B — Bold rewrite with stronger structural choices...</div>
    </div>
    """

    st.markdown(f"""
    <div class="upgrade-wall">
      <div class="upgrade-blur">{locked_preview_html}</div>
    </div>
    """, unsafe_allow_html=True)


def show_upgrade_card():
    """Inline upgrade CTA with coupon — shown after the teaser."""
    price_display  = st.secrets.get("PRICE_DISPLAY", "$9")
    price_subtitle = st.secrets.get("PRICE_SUBTITLE", "one-time · unlimited audits")
    payment_link   = st.secrets.get("STRIPE_PAYMENT_LINK", "")

    st.markdown(f"""
    <div class="upgrade-card">
      <div class="uc-eyebrow">Unlock Your Full Report</div>
      <h3>See exactly what to fix — and get it rewritten for you</h3>
      <p class="uc-sub">Your full report includes everything below, ready to use immediately.</p>
      <ul>
        <li>All 6 category analyses with specific fixes</li>
        <li>Your top 3 priority changes (ranked by impact)</li>
        <li>Two complete rewritten blurbs — copy straight to your retailer page</li>
        <li>Trope check table (if you provided tropes)</li>
        <li>Downloadable Word document</li>
        <li>Unlimited audits — run it as many times as you want</li>
      </ul>
      <div class="price-badge">{price_display}</div>
      <div style="color:#888;font-size:0.85rem;margin-bottom:1.2rem;">{price_subtitle}</div>
    </div>
    """, unsafe_allow_html=True)

    if payment_link:
        st.link_button(
            f"Unlock Full Report — {price_display} →",
            payment_link,
            use_container_width=True,
        )

    st.markdown('<div class="divider-or">or</div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="coupon-section">
      <div class="label">🎓 Writing Wives Skool Member? Enter your coupon for free access.</div>
    </div>
    """, unsafe_allow_html=True)

    with st.form("coupon_form"):
        code  = st.text_input("Coupon code", placeholder="", label_visibility="collapsed")
        apply = st.form_submit_button("Apply Coupon →")

    if apply:
        if code and check_coupon(code):
            grant_access(reason="coupon")
            st.rerun()
        else:
            st.error("That coupon code isn't valid. Double-check and try again.")


def render_ad_package(ad_data, book_title):
    """Render the full FB Ad Package report."""
    st.markdown(f"""
    <div class="ad-section-header">
      <div class="ad-eyebrow">FB &amp; Instagram Ad Copy</div>
      <h2>Your Ad Package — {book_title or 'Ready to Use'}</h2>
    </div>
    """, unsafe_allow_html=True)

    best = ad_data.get("best_picks", {})
    best_hl  = best.get("headline", 0)
    best_pt  = best.get("primary_text", 1)
    best_desc = best.get("description", 0)

    # ── Headlines ────────────────────────────────────────────────────────────
    st.markdown('<div class="section-h2">📰 Headlines</div>', unsafe_allow_html=True)
    st.caption("Under 40 characters each. Copy straight into the Headline field in Ads Manager.")
    for i, hl in enumerate(ad_data.get("headlines", [])):
        is_best = (i == best_hl)
        best_badge = '<div class="best-pick-badge">⭐ Best Pick</div>' if is_best else ""
        card_class = "headline-card best" if is_best else "headline-card"
        st.markdown(f"""
        <div class="{card_class}">
          {best_badge}
          <div class="hl-text">{hl.get('text','')}</div>
          <div class="hl-meta">{hl.get('chars', len(hl.get('text','')))} chars &nbsp;·&nbsp; {hl.get('angle','')}</div>
        </div>
        """, unsafe_allow_html=True)

    # ── Primary Texts ─────────────────────────────────────────────────────────
    st.markdown('<div class="section-h2">📝 Primary Text Options</div>', unsafe_allow_html=True)
    st.caption("Three length variations — test them or pick the one that fits your campaign.")
    for i, pt in enumerate(ad_data.get("primary_texts", [])):
        is_best = (i == best_pt)
        best_badge = '<div class="best-pick-badge">⭐ Best Pick</div>' if is_best else ""
        card_class = "pt-card best" if is_best else "pt-card"
        st.markdown(f"""
        <div class="{card_class}">
          {best_badge}
          <div class="pt-label">{pt.get('label','')}</div>
          <div class="pt-text">{pt.get('text','')}</div>
          <div class="pt-notes">💡 {pt.get('notes','')}</div>
        </div>
        """, unsafe_allow_html=True)

    # ── Descriptions ──────────────────────────────────────────────────────────
    st.markdown('<div class="section-h2">🏷️ Description Field Options</div>', unsafe_allow_html=True)
    st.caption("The small line under the headline in the ad. Two options — pick one or A/B test.")
    for i, desc in enumerate(ad_data.get("descriptions", [])):
        is_best = (i == best_desc)
        best_badge = '<div class="best-pick-badge">⭐ Best Pick</div>' if is_best else ""
        card_class = "desc-card best" if is_best else "desc-card"
        st.markdown(f"""
        <div class="{card_class}">
          {best_badge}
          <div class="desc-text">{desc.get('text','')}</div>
          <div class="desc-notes">{desc.get('notes','')}</div>
        </div>
        """, unsafe_allow_html=True)

    # ── Targeting Note ────────────────────────────────────────────────────────
    targeting = ad_data.get("targeting_note")
    if targeting:
        st.markdown('<div class="section-h2">🎯 Audience Targeting Suggestion</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="targeting-box">🎯 {targeting}</div>', unsafe_allow_html=True)

    # ── Downloads ─────────────────────────────────────────────────────────────
    st.divider()
    st.markdown("**Download your ad package:**")
    col_a, col_b = st.columns(2)
    with col_a:
        st.download_button(
            "⬇ Download as JSON",
            data=json.dumps(ad_data, indent=2, ensure_ascii=False).encode("utf-8"),
            file_name=f"fb-ad-package-{(book_title or 'report').replace(' ','-').lower()}.json",
            mime="application/json",
            key="dl_ad_json",
        )
    with col_b:
        if DOCX_AVAILABLE:
            st.download_button(
                "⬇ Download as Word Doc",
                data=generate_ad_docx(ad_data, book_title),
                file_name=f"FB Ad Package - {book_title or 'Report'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_ad_docx",
            )


def show_ad_upgrade_card():
    """Upsell card shown at the bottom of the full blurb audit report."""
    ad_price   = st.secrets.get("AD_PRICE_DISPLAY", "$25")
    ad_sub     = st.secrets.get("AD_PRICE_SUBTITLE", "one-time · ad copy for this blurb")
    ad_link    = st.secrets.get("STRIPE_AD_LINK", "")

    st.markdown(f"""
    <div class="ad-upgrade-card">
      <div class="uc-eyebrow">Ready to Run Ads on This Book?</div>
      <h3>Get your complete FB &amp; Instagram Ad Package</h3>
      <p class="uc-sub">Your audit identified what's strong in this blurb — now let's turn it into ads that sell.</p>
      <ul>
        <li>5 headline variations (under 40 chars, tested angles)</li>
        <li>3 primary text options — short, medium, long</li>
        <li>2 description field options</li>
        <li>⭐ Best Pick recommendations for each element</li>
        <li>Audience targeting suggestion based on your genre</li>
        <li>Downloadable Word doc — copy straight into Ads Manager</li>
      </ul>
      <div class="ad-price-badge">{ad_price}</div>
      <div style="color:#bbb;font-size:0.85rem;margin-bottom:1.2rem;">{ad_sub}</div>
    </div>
    """, unsafe_allow_html=True)

    if ad_link:
        st.link_button(
            f"Get My Ad Package — {ad_price} →",
            ad_link,
            use_container_width=True,
        )

    st.markdown('<div class="divider-or">or</div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="coupon-section">
      <div class="label">🎓 Writing Wives Skool Member? Enter your coupon for free access.</div>
    </div>
    """, unsafe_allow_html=True)

    with st.form("ad_coupon_form"):
        code  = st.text_input("Coupon code", placeholder="", label_visibility="collapsed", key="ad_coupon_input")
        apply = st.form_submit_button("Apply Coupon →")

    if apply:
        if code and check_ad_coupon(code):
            grant_ad_access(reason="coupon")
            st.rerun()
        else:
            st.error("That coupon code isn't valid. Double-check and try again.")


# ── Word doc generation ───────────────────────────────────────────────────────

def generate_docx(data, blurb, book_title):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    h = doc.add_heading('Blurb Audit Report', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if book_title:
        p = doc.add_paragraph(book_title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(13)
        p.runs[0].font.color.rgb = RGBColor(0x6B, 0x2D, 0x5E)
    doc.add_paragraph("")
    doc.add_heading("Blurb as Submitted", 1)
    doc.add_paragraph(blurb).runs[0].font.italic = True
    doc.add_heading("Audit Summary", 1)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = "Category"
    t.rows[0].cells[1].text = "Rating"
    for key, label in CATEGORY_LABELS.items():
        r = t.add_row().cells
        r[0].text = label
        r[1].text = data.get("summary", {}).get(key, "")
    doc.add_heading("Detailed Findings", 1)
    for key, label in CATEGORY_LABELS.items():
        cat = data.get("categories", {}).get(key, {})
        if not cat: continue
        doc.add_heading(f"{label}  {cat.get('rating','')}", 2)
        for obs in cat.get("observations", []):
            doc.add_paragraph(obs, style='List Bullet')
        if cat.get("suggested_fix"):
            p = doc.add_paragraph(f"Suggested fix: {cat['suggested_fix']}")
            p.runs[0].font.italic = True
    if data.get("trope_check"):
        doc.add_heading("Trope Check", 1)
        t2 = doc.add_table(rows=1, cols=3)
        t2.style = 'Table Grid'
        t2.rows[0].cells[0].text = "Trope"
        t2.rows[0].cells[1].text = "Status"
        t2.rows[0].cells[2].text = "Evidence / Suggestion"
        for tr in data["trope_check"]:
            r = t2.add_row().cells
            r[0].text = tr.get("trope", "")
            r[1].text = tr.get("status", "")
            r[2].text = tr.get("evidence", "")
    doc.add_heading("Top 3 Priority Fixes", 1)
    for fix in data.get("priority_fixes", []):
        doc.add_paragraph(fix, style='List Number')
    doc.add_page_break()
    doc.add_heading("Rewritten Blurb Options", 1)
    for key in ("rewrite_a", "rewrite_b"):
        rw = data.get(key, {})
        if not rw: continue
        doc.add_heading(rw.get("label", key), 2)
        doc.add_paragraph(rw.get("text", "")).runs[0].font.italic = True
        if rw.get("notes"):
            p = doc.add_paragraph(f"Note: {rw['notes']}")
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_ad_docx(ad_data, book_title):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    h = doc.add_heading('Facebook Ad Package', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if book_title:
        p = doc.add_paragraph(book_title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(13)
        p.runs[0].font.bold = True

    best = ad_data.get("best_picks", {})
    doc.add_paragraph("")

    # Headlines
    doc.add_heading("Headlines (copy into Ads Manager Headline field)", 1)
    doc.add_paragraph("Under 40 characters each. Vary by angle for A/B testing.").runs[0].font.italic = True
    for i, hl in enumerate(ad_data.get("headlines", [])):
        is_best = (i == best.get("headline", 0))
        label = f"{'⭐ BEST PICK — ' if is_best else ''}Headline {i+1}: {hl.get('text','')}"
        p = doc.add_paragraph(label, style='List Number')
        if is_best:
            p.runs[0].font.bold = True
        doc.add_paragraph(f"   {hl.get('chars','')} chars · {hl.get('angle','')}").runs[0].font.size = Pt(9)

    # Primary Texts
    doc.add_heading("Primary Text Options", 1)
    for i, pt in enumerate(ad_data.get("primary_texts", [])):
        is_best = (i == best.get("primary_text", 1))
        heading_text = f"{'⭐ BEST PICK — ' if is_best else ''}{pt.get('label','')}"
        h2 = doc.add_heading(heading_text, 2)
        if is_best:
            for run in h2.runs:
                run.font.bold = True
        doc.add_paragraph(pt.get("text", "")).runs[0].font.italic = False
        if pt.get("notes"):
            p = doc.add_paragraph(f"💡 Use when: {pt['notes']}")
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Descriptions
    doc.add_heading("Description Field Options", 1)
    doc.add_paragraph("The short line below the headline in the ad.").runs[0].font.italic = True
    for i, desc in enumerate(ad_data.get("descriptions", [])):
        is_best = (i == best.get("description", 0))
        label = f"{'⭐ BEST PICK — ' if is_best else ''}Option {i+1}: {desc.get('text','')}"
        p = doc.add_paragraph(label, style='List Number')
        if is_best:
            p.runs[0].font.bold = True
        if desc.get("notes"):
            doc.add_paragraph(f"   {desc['notes']}").runs[0].font.size = Pt(9)

    # Targeting
    targeting = ad_data.get("targeting_note")
    if targeting:
        doc.add_heading("Audience Targeting Suggestion", 1)
        doc.add_paragraph(targeting)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═════════════════════════════════════════════════════════════════════════════
# MAIN APP
# ═════════════════════════════════════════════════════════════════════════════

# Back to home
if st.button("← Back to All Tools", key="back_home"):
    st.switch_page("home.py")

# Header
_logo_b64 = get_logo_b64()
_logo_html = f'<div class="logo-wrap"><img src="data:image/png;base64,{_logo_b64}" alt="The Writing Wives"></div>' if _logo_b64 else ""
st.markdown(f"""
<div class="header-block">
  {_logo_html}
  <div class="eyebrow">Genre &amp; Trope AI Classroom</div>
  <h1>Blurb Auditor</h1>
  <p class="sub">Paste your blurb. Get an instant scored audit + two rewritten versions — ready to use.</p>
</div>
""", unsafe_allow_html=True)

# ── Check Stripe session_id in URL (blurb audit) ──────────────────────────────
if not is_authenticated():
    session_id = st.query_params.get("session_id")
    if session_id and not st.session_state.get("stripe_checked"):
        st.session_state["stripe_checked"] = True
        with st.spinner("Verifying your payment..."):
            ok, email = verify_stripe_session(session_id)
        if ok:
            grant_access(reason="stripe", email=email)
            st.rerun()
        else:
            st.warning("We couldn't verify that payment yet — it may still be processing. Try refreshing in a moment.")

# ── Check Stripe ad_session_id in URL (ad package) ────────────────────────────
if not is_ad_authenticated():
    ad_session_id = st.query_params.get("ad_session_id")
    if ad_session_id and not st.session_state.get("ad_stripe_checked"):
        st.session_state["ad_stripe_checked"] = True
        with st.spinner("Verifying your ad package payment..."):
            ok, email = verify_stripe_session(ad_session_id)
        if ok:
            grant_ad_access(reason="stripe", email=email)
            st.rerun()
        else:
            st.warning("We couldn't verify that payment yet — it may still be processing. Try refreshing in a moment.")

# ── Access granted banners ────────────────────────────────────────────────────
if is_authenticated():
    reason = st.session_state.get("access_reason", "")
    email  = st.session_state.get("access_email")
    if reason == "stripe":
        greeting = f"Payment confirmed{f' for {email}' if email else ''}."
        st.markdown(f"""
        <div class="success-banner">✅ <strong>{greeting}</strong> Full access unlocked.</div>
        <div class="bookmark-tip">🔖 <strong>Bookmark this page</strong> — your URL contains your access token. Come back any time for more audits.</div>
        """, unsafe_allow_html=True)
    elif reason == "coupon":
        st.markdown('<div class="success-banner">🎓 <strong>Member access granted.</strong> Run as many audits as you need.</div>', unsafe_allow_html=True)

# ── Audit form ────────────────────────────────────────────────────────────────
with st.form("audit_form"):
    book_title = st.text_input("Book Title", placeholder="")
    col1, col2 = st.columns(2)
    with col1:
        genre = st.text_input("Genre", placeholder="")
    with col2:
        tense = st.selectbox(
            "POV / Tense",
            ["First person present", "First person past", "Third person present",
             "Third person past", "Second person", "Not sure"],
        )
    tropes = st.text_input(
        "Your top tropes (optional but recommended)",
        placeholder="e.g. Chosen One, Found Family, Race Against Time",
    )
    blurb = st.text_area(
        "Your Blurb",
        height=200,
        placeholder="Paste your full blurb here exactly as it appears on your retailer page...",
    )
    submit_label = "Audit My Blurb →" if is_authenticated() else "Get My Free Preview →"
    submitted = st.form_submit_button(submit_label)

# ── On submit ─────────────────────────────────────────────────────────────────
if submitted:
    if not blurb.strip():
        st.warning("Please paste your blurb before submitting.")
    else:
        spinner_msg = "Auditing your blurb..." if is_authenticated() else "Analysing your blurb — this takes about 15 seconds..."
        with st.spinner(spinner_msg):
            try:
                result = call_openrouter(blurb, genre, tense, tropes, book_title)
                # Store in session so coupon unlock can show it immediately
                st.session_state["last_audit"] = {
                    "result": result, "blurb": blurb,
                    "book_title": book_title, "genre": genre
                }
            except json.JSONDecodeError as e:
                st.error("The AI returned an unexpected format. Please try again.")
                st.exception(e)
                st.stop()
            except Exception as e:
                st.error("Something went wrong. Check your API key and try again.")
                st.exception(e)
                st.stop()

        st.success("Done!" if is_authenticated() else "Preview ready!")
        st.divider()

        if is_authenticated():
            render_full_audit(result, blurb, book_title)
            # ── Ad Package upsell / render ─────────────────────────────────
            st.divider()
            if is_ad_authenticated():
                with st.spinner("Generating your FB Ad Package..."):
                    try:
                        ad_result = call_openrouter_ads(blurb, genre, book_title, audit_data=result)
                        st.session_state["last_ad_audit"] = {"result": ad_result, "book_title": book_title}
                    except Exception as e:
                        st.error("Couldn't generate the ad package. Please try again.")
                        st.exception(e)
                        ad_result = None
                if ad_result:
                    render_ad_package(ad_result, book_title)
            else:
                show_ad_upgrade_card()
        else:
            render_teaser_audit(result, blurb)
            show_upgrade_card()

# ── If they just unlocked via coupon and a prior audit is stored ──────────────
elif is_authenticated() and st.session_state.get("last_audit") and st.session_state.get("access_reason") == "coupon":
    prior = st.session_state["last_audit"]
    st.success("✅ Access granted! Here's your full report:")
    st.divider()
    render_full_audit(prior["result"], prior["blurb"], prior["book_title"])
    # Show ad upsell or render after coupon unlock too
    st.divider()
    if is_ad_authenticated():
        with st.spinner("Generating your FB Ad Package..."):
            try:
                ad_result = call_openrouter_ads(prior["blurb"], prior.get("genre",""), prior["book_title"], audit_data=prior["result"])
                st.session_state["last_ad_audit"] = {"result": ad_result, "book_title": prior["book_title"]}
            except Exception:
                ad_result = None
        if ad_result:
            render_ad_package(ad_result, prior["book_title"])
    else:
        show_ad_upgrade_card()
    # Clear so it doesn't re-show on every page reload
    del st.session_state["last_audit"]

# ── If they just unlocked ad package via coupon ───────────────────────────────
elif is_authenticated() and is_ad_authenticated() and st.session_state.get("ad_access_reason") == "coupon":
    if st.session_state.get("last_audit"):
        prior = st.session_state["last_audit"]
        st.success("🎉 Ad package access granted!")
        with st.spinner("Generating your FB Ad Package..."):
            try:
                ad_result = call_openrouter_ads(prior["blurb"], prior.get("genre",""), prior["book_title"], audit_data=prior["result"])
            except Exception:
                ad_result = None
        if ad_result:
            render_ad_package(ad_result, prior["book_title"])

# Footer
st.markdown("""
<div class="footer-note">
  The Writing Wives · Genre &amp; Trope AI Classroom ·
  <a href="https://thewritingwives.com" style="color:#999;">thewritingwives.com</a>
</div>
""", unsafe_allow_html=True)
