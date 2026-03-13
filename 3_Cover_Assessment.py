import streamlit as st
from openai import OpenAI
import json, re, io, base64
from pathlib import Path

from affiliate_utils import capture_ref, affiliate_payment_link
capture_ref()

try:
    import stripe as stripe_lib
    STRIPE_AVAILABLE = True
except ImportError:
    STRIPE_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False



PRIMARY    = "#1A1A1A"
PRIMARY_MID= "#2E2E2E"
GOLD       = "#D4B36E"
GOLD_DARK  = "#A8863A"
GOLD_LIGHT = "#FAF4E6"
GOLD_MID   = "#F5EDD6"
GOLD_BG    = "#FFFBEC"
GREEN      = "#1A7A4A"
GREEN_BG   = "#D4EDDA"

CATEGORY_LABELS = {
    "genre_signal":       "1. Genre Signal",
    "thumbnail_performance": "2. Thumbnail Performance",
    "title_readability":  "3. Title Readability",
    "professional_quality": "4. Professional Quality",
    "market_positioning": "5. Market Positioning",
    "color_mood_match":   "6. Colour & Mood Match",
}

def get_logo_b64():
    p = Path(__file__).parent / "logo.png"
    if p.exists():
        with open(p, "rb") as f:
            return base64.b64encode(f.read()).decode()
    return None

def get_secret(key, default=""):
    try:
        return st.secrets.get(key, default)
    except Exception:
        return default

def verify_stripe_session(session_id):
    if not STRIPE_AVAILABLE:
        return False, None
    try:
        stripe_lib.api_key = get_secret("STRIPE_SECRET_KEY")
        session = stripe_lib.checkout.Session.retrieve(session_id)
        if session.payment_status == "paid":
            email = session.customer_details.email if session.customer_details else None
            return True, email
    except Exception:
        pass
    return False, None

def is_authenticated():
    return (st.session_state.get("cover_access_granted", False) or
            st.session_state.get("lifetime_access", False))

def grant_access(reason="", email=None):
    st.session_state["cover_access_granted"] = True
    st.session_state["cover_access_reason"]  = reason
    if email:
        st.session_state["cover_access_email"] = email

def check_coupon(code):
    raw = get_secret("COVER_COUPON_CODES", get_secret("COUPON_CODES", ""))
    valid = {c.strip().upper() for c in raw.split(",") if c.strip()}
    return code.strip().upper() in valid

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
  html, body, [class*="css"] {{ font-family: 'Segoe UI', Arial, sans-serif; }}
  .block-container {{ max-width: 760px; padding-top: 1.5rem; }}
  .header-block {{ background:{PRIMARY}; padding:2rem 2.5rem 1.6rem; border-radius:12px; margin-bottom:2rem; border-bottom:3px solid {GOLD}; }}
  .header-block .logo-wrap {{ text-align:center; margin-bottom:1.1rem; }}
  .header-block .logo-wrap img {{ height:140px; max-width:100%; object-fit:contain; }}
  .header-block .eyebrow {{ color:{GOLD}; font-size:0.72rem; font-weight:700; letter-spacing:0.18em; text-transform:uppercase; margin:0 0 0.3rem; text-align:center; }}
  .header-block h1 {{ color:#fff; font-size:2rem; font-weight:900; margin:0 0 0.35rem; line-height:1.15; text-align:center; }}
  .header-block .sub {{ color:#aaa; font-size:0.95rem; margin:0; text-align:center; }}
  .r-strong  {{ background:#D4EDDA; color:#155724; padding:3px 12px; border-radius:20px; font-weight:700; font-size:0.88rem; display:inline-block; }}
  .r-warning {{ background:#FFF3CD; color:#856404; padding:3px 12px; border-radius:20px; font-weight:700; font-size:0.88rem; display:inline-block; }}
  .r-missing {{ background:#F8D7DA; color:#721C24; padding:3px 12px; border-radius:20px; font-weight:700; font-size:0.88rem; display:inline-block; }}
  .section-h2 {{ color:{PRIMARY}; font-size:1.3rem; font-weight:900; border-bottom:3px solid {GOLD}; padding-bottom:6px; margin:2rem 0 1rem; }}
  .cat-heading {{ color:{PRIMARY}; font-size:1.1rem; font-weight:700; border-bottom:2px solid {GOLD}; padding-bottom:4px; margin:1.6rem 0 0.6rem; }}
  .fix-box {{ background:{GOLD_LIGHT}; border-left:4px solid {GOLD}; padding:0.75rem 1rem; border-radius:0 8px 8px 0; margin-top:0.5rem; font-size:0.92rem; line-height:1.65; }}
  .priority-box {{ background:{GOLD_BG}; border-left:5px solid {GOLD_DARK}; padding:0.9rem 1.1rem; border-radius:0 10px 10px 0; margin:0.4rem 0; font-size:0.93rem; line-height:1.6; }}
  .sum-grid {{ display:grid; grid-template-columns:1fr 1fr; gap:8px; margin:0.75rem 0 1.5rem; }}
  .sum-cell {{ background:{GOLD_MID}; border-radius:8px; padding:10px 14px; display:flex; align-items:center; justify-content:space-between; font-size:0.9rem; border:1px solid #E8D9B0; }}
  .sum-cell .sum-label {{ font-weight:600; color:#333; }}
  .score-big {{ text-align:center; padding:1.5rem; background:{PRIMARY}; border-radius:12px; margin:1rem 0; border-bottom:3px solid {GOLD}; }}
  .score-big .score-num {{ font-size:3.5rem; font-weight:900; color:{GOLD}; line-height:1; }}
  .score-big .score-label {{ color:#aaa; font-size:0.9rem; margin-top:0.3rem; }}
  .score-big .score-status {{ font-size:1rem; font-weight:700; margin-top:0.5rem; }}
  .hire-box {{ background:#FFF4E5; border-left:4px solid #F6A623; border-radius:0 8px 8px 0; padding:0.85rem 1.1rem; margin:0.5rem 0; font-size:0.92rem; color:#7A4100; }}
  .comp-box {{ background:{GOLD_LIGHT}; border-left:4px solid {GOLD}; border-radius:0 8px 8px 0; padding:0.85rem 1.1rem; margin:0.5rem 0; font-size:0.92rem; color:{PRIMARY}; }}
  .upgrade-card {{ background:#fff; border:2px solid {GOLD_MID}; border-top:4px solid {GOLD}; border-radius:14px; padding:2rem; margin:1.5rem 0; text-align:center; }}
  .upgrade-card .uc-eyebrow {{ font-size:0.75rem; font-weight:700; letter-spacing:0.15em; text-transform:uppercase; color:{GOLD_DARK}; margin-bottom:0.4rem; }}
  .upgrade-card h3 {{ font-size:1.4rem; font-weight:900; color:{PRIMARY}; margin:0 0 0.3rem; }}
  .upgrade-card .uc-sub {{ color:#666; font-size:0.92rem; margin-bottom:1.2rem; }}
  .upgrade-card ul {{ text-align:left; display:inline-block; list-style:none; padding:0; margin:0 auto 1.5rem; }}
  .upgrade-card ul li {{ padding:3px 0; font-size:0.92rem; color:#333; }}
  .upgrade-card ul li::before {{ content:"✓ "; color:{GOLD_DARK}; font-weight:700; }}
  .price-badge {{ display:inline-block; background:{PRIMARY}; color:{GOLD}; font-size:1.6rem; font-weight:900; padding:6px 28px; border-radius:40px; margin-bottom:1.2rem; border:2px solid {GOLD}; }}
  .divider-or {{ display:flex; align-items:center; gap:12px; color:#bbb; font-size:0.85rem; margin:1.2rem 0; }}
  .divider-or::before, .divider-or::after {{ content:""; flex:1; border-top:1px solid #e0e0e0; }}
  .coupon-section {{ background:{GOLD_BG}; border:1px solid #e8d78a; border-left:4px solid {GOLD}; border-radius:0 8px 8px 0; padding:1rem 1.25rem; margin-bottom:1rem; }}
  .coupon-section .label {{ font-weight:700; font-size:0.9rem; color:#5C4A00; margin-bottom:0.4rem; }}
  .success-banner {{ background:{GREEN_BG}; border:1px solid #b8ddc8; border-left:4px solid {GREEN}; border-radius:0 8px 8px 0; padding:1rem 1.25rem; margin-bottom:1.5rem; font-size:0.92rem; color:#155724; }}
  .bookmark-tip {{ background:#EEF4FF; border:1px solid #c5d8f8; border-left:4px solid #3B6FD4; border-radius:0 8px 8px 0; padding:0.85rem 1.1rem; margin-bottom:1.5rem; font-size:0.88rem; color:#1a3a6b; }}
  .stButton > button {{ background:{PRIMARY} !important; color:{GOLD} !important; font-weight:700 !important; padding:0.65rem 2rem !important; border-radius:8px !important; border:2px solid {GOLD} !important; font-size:1rem !important; width:100%; }}
  .stButton > button:hover {{ background:{GOLD} !important; color:{PRIMARY} !important; }}
  .stLinkButton > a {{ background:{PRIMARY} !important; color:{GOLD} !important; font-weight:700 !important; border-radius:8px !important; font-size:1rem !important; border:2px solid {GOLD} !important; }}
  .footer-note {{ text-align:center; color:#aaa; font-size:0.78rem; margin-top:3rem; padding-top:1rem; border-top:1px solid {GOLD_MID}; }}
</style>
""", unsafe_allow_html=True)

# ── AI Prompt ─────────────────────────────────────────────────────────────────
COVER_SYSTEM_PROMPT = """You are the Cover Assessment AI for The Writing Wives Genre & Trope AI Classroom.

Analyse this book cover image from a professional book marketing perspective. Be honest — authors need truth, not flattery.

Return ONLY valid JSON. No markdown, no code fences. Exact structure:

{
  "summary": {
    "genre_signal": "✅ Strong",
    "thumbnail_performance": "⚠️ Needs Work",
    "title_readability": "✅ Strong",
    "professional_quality": "⚠️ Needs Work",
    "market_positioning": "❌ Missing",
    "color_mood_match": "✅ Strong"
  },
  "categories": {
    "genre_signal":          {"rating": "✅", "observations": ["..."], "suggested_fix": null},
    "thumbnail_performance": {"rating": "⚠️", "observations": ["..."], "suggested_fix": "..."},
    "title_readability":     {"rating": "✅", "observations": ["..."], "suggested_fix": null},
    "professional_quality":  {"rating": "⚠️", "observations": ["..."], "suggested_fix": "..."},
    "market_positioning":    {"rating": "❌", "observations": ["..."], "suggested_fix": "..."},
    "color_mood_match":      {"rating": "✅", "observations": ["..."], "suggested_fix": null}
  },
  "genre_detected": "What genre this cover signals visually (may differ from the author's stated genre)",
  "overall_score": 72,
  "priority_fixes": ["Fix 1 — specific, actionable.", "Fix 2 — specific.", "Fix 3 — specific."],
  "comparable_style": "Brief description of what top-performing covers in this genre look like.",
  "hire_note": "Honest 1–2 sentence assessment: is this a full redesign situation or targeted refinements?"
}

Categories to assess:
- genre_signal: Does this cover immediately scream its genre? Would a reader scrolling Amazon know what they're getting?
- thumbnail_performance: At 160×260px (Amazon thumbnail), is the cover still clear and compelling?
- title_readability: Is the title immediately legible? Font size, contrast, placement, and style.
- professional_quality: Does it look indie-typical or does it compete with Big 5 covers?
- market_positioning: Does it fit among current bestsellers in its genre/sub-genre?
- color_mood_match: Does the colour palette emotionally match the genre expectations?

overall_score: 0–100. 80+ is market-ready. 60–79 is functional but needs work. Below 60 needs significant redesign.
If the visual genre signal doesn't match the author's stated genre, flag this as a critical issue."""

def call_cover_assessment(image_b64, genre, book_title, series="", mime_type="image/jpeg"):
    api_key = get_secret("OPENROUTER_API_KEY")
    if not api_key:
        st.error("API key not configured.")
        st.stop()
    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
    model  = get_secret("MODEL", "anthropic/claude-sonnet-4-5")
    user_text = f"Please assess this book cover."
    if book_title: user_text += f"\nBook title: {book_title}"
    if series:     user_text += f"\nSeries: {series}"
    if genre:      user_text += f"\nAuthor's intended genre: {genre}"
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": COVER_SYSTEM_PROMPT},
            {"role": "user", "content": [
                {"type": "text", "text": user_text},
                {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_b64}"}},
            ]},
        ],
        temperature=0.4,
        max_tokens=3000,
    )
    raw = response.choices[0].message.content.strip()
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    return json.loads(raw.strip())

def rating_badge(r):
    if "✅" in r:   return '<span class="r-strong">✅ Strong</span>'
    elif "⚠️" in r: return '<span class="r-warning">⚠️ Needs Work</span>'
    else:            return '<span class="r-missing">❌ Missing</span>'

def render_full_assessment(data, book_title):
    # Overall score
    score  = data.get("overall_score", 0)
    genre_detected = data.get("genre_detected", "")
    if score >= 80:
        status_html = f'<div class="score-status" style="color:{GOLD};">✅ Market-Ready</div>'
    elif score >= 60:
        status_html = f'<div class="score-status" style="color:#F6A623;">⚠️ Functional — Needs Refinements</div>'
    else:
        status_html = f'<div class="score-status" style="color:#E74C3C;">❌ Needs Significant Redesign</div>'

    st.markdown(f"""
    <div class="score-big">
      <div class="score-num">{score}</div>
      <div class="score-label">Cover Score out of 100</div>
      {status_html}
    </div>
    """, unsafe_allow_html=True)

    if genre_detected:
        st.markdown(f'<div class="comp-box">🔍 <strong>Genre signal detected:</strong> {genre_detected}</div>', unsafe_allow_html=True)

    # Summary grid
    st.markdown('<div class="section-h2">Assessment Summary</div>', unsafe_allow_html=True)
    summary = data.get("summary", {})
    grid = '<div class="sum-grid">'
    for key, label in CATEGORY_LABELS.items():
        grid += f'<div class="sum-cell"><span class="sum-label">{label}</span>{rating_badge(summary.get(key,""))}</div>'
    grid += '</div>'
    st.markdown(grid, unsafe_allow_html=True)

    # Detailed categories
    st.markdown('<div class="section-h2">Detailed Findings</div>', unsafe_allow_html=True)
    cats = data.get("categories", {})
    for key, label in CATEGORY_LABELS.items():
        cat = cats.get(key, {})
        if not cat:
            continue
        st.markdown(f'<div class="cat-heading">{label} &nbsp; {rating_badge(cat.get("rating",""))}</div>', unsafe_allow_html=True)
        for obs in cat.get("observations", []):
            st.markdown(f"- {obs}")
        if cat.get("suggested_fix"):
            st.markdown(f'<div class="fix-box"><strong>Suggested fix:</strong> {cat["suggested_fix"]}</div>', unsafe_allow_html=True)

    # Priority fixes
    st.markdown('<div class="section-h2">Top 3 Priority Fixes</div>', unsafe_allow_html=True)
    for fix in data.get("priority_fixes", []):
        st.markdown(f'<div class="priority-box">{fix}</div>', unsafe_allow_html=True)

    # Comparable style
    if data.get("comparable_style"):
        st.markdown('<div class="section-h2">What Top Covers in This Genre Look Like</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="comp-box">🎨 {data["comparable_style"]}</div>', unsafe_allow_html=True)

    # Hire note
    if data.get("hire_note"):
        st.markdown('<div class="section-h2">Designer Recommendation</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="hire-box">💡 {data["hire_note"]}</div>', unsafe_allow_html=True)

    # Downloads
    st.divider()
    st.markdown("**Download your assessment:**")
    ca, cb = st.columns(2)
    with ca:
        st.download_button("⬇ Download as JSON",
            data=json.dumps(data, indent=2, ensure_ascii=False).encode(),
            file_name=f"cover-assessment-{(book_title or 'report').replace(' ','-').lower()}.json",
            mime="application/json", key="dl_cover_json")
    with cb:
        if DOCX_AVAILABLE:
            st.download_button("⬇ Download as Word Doc",
                data=generate_cover_docx(data, book_title),
                file_name=f"Cover Assessment - {book_title or 'Report'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_cover_docx")

def render_teaser(data):
    score = data.get("overall_score", 0)
    summary = data.get("summary", {})
    issue_count = sum(1 for v in summary.values() if "⚠️" in v or "❌" in v)
    st.markdown(f"""
    <div style="background:{GOLD_LIGHT};border:2px solid {GOLD_MID};border-left:5px solid {GOLD};border-radius:0 10px 10px 0;padding:1.1rem 1.3rem;margin:1rem 0;">
      <strong>Your cover scored {score}/100</strong> — with {issue_count} area{'s' if issue_count != 1 else ''} flagged.
      Unlock your full report to see exactly what to fix.
    </div>
    """, unsafe_allow_html=True)
    # Summary grid still visible
    st.markdown('<div class="section-h2">Assessment Summary</div>', unsafe_allow_html=True)
    grid = '<div class="sum-grid">'
    for key, label in CATEGORY_LABELS.items():
        grid += f'<div class="sum-cell"><span class="sum-label">{label}</span>{rating_badge(summary.get(key,""))}</div>'
    grid += '</div>'
    st.markdown(grid, unsafe_allow_html=True)

def generate_cover_docx(data, book_title):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    h = doc.add_heading('Cover Assessment Report', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if book_title:
        p = doc.add_paragraph(book_title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(13)
        p.runs[0].font.bold = True
    doc.add_paragraph(f"Overall Score: {data.get('overall_score', 0)}/100")
    if data.get("genre_detected"):
        doc.add_paragraph(f"Genre detected: {data['genre_detected']}")
    doc.add_heading("Assessment Summary", 1)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = "Category"
    t.rows[0].cells[1].text = "Rating"
    for key, label in CATEGORY_LABELS.items():
        r = t.add_row().cells
        r[0].text = label
        r[1].text = data.get("summary", {}).get(key, "")
    doc.add_heading("Detailed Findings", 1)
    for key, label in CATEGORY_LABELS.items():
        cat = data.get("categories", {}).get(key, {})
        if not cat: continue
        doc.add_heading(f"{label}  {cat.get('rating','')}", 2)
        for obs in cat.get("observations", []):
            doc.add_paragraph(obs, style='List Bullet')
        if cat.get("suggested_fix"):
            p = doc.add_paragraph(f"Fix: {cat['suggested_fix']}")
            p.runs[0].font.italic = True
    doc.add_heading("Top 3 Priority Fixes", 1)
    for fix in data.get("priority_fixes", []):
        doc.add_paragraph(fix, style='List Number')
    if data.get("comparable_style"):
        doc.add_heading("What Top Covers in This Genre Look Like", 1)
        doc.add_paragraph(data["comparable_style"])
    if data.get("hire_note"):
        doc.add_heading("Designer Recommendation", 1)
        doc.add_paragraph(data["hire_note"])
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def show_upgrade_card():
    price = get_secret("COVER_PRICE_DISPLAY", "$19")
    sub   = get_secret("COVER_PRICE_SUBTITLE", "one-time · per cover")
    link  = affiliate_payment_link(get_secret("STRIPE_COVER_LINK", ""))
    st.markdown(f"""
    <div class="upgrade-card">
      <div class="uc-eyebrow">Unlock Your Full Assessment</div>
      <h3>See exactly what's working and what needs to change</h3>
      <p class="uc-sub">You've seen your scores — now get the full breakdown with specific, actionable fixes.</p>
      <ul>
        <li>6-category detailed findings</li>
        <li>Specific fix for every ⚠️ or ❌</li>
        <li>Top 3 priority changes (ranked by impact)</li>
        <li>What top covers in your genre look like</li>
        <li>Honest designer recommendation</li>
        <li>Downloadable Word doc</li>
      </ul>
      <div class="price-badge">{price}</div>
      <div style="color:#888;font-size:0.85rem;margin-bottom:1.2rem;">{sub}</div>
    </div>
    """, unsafe_allow_html=True)
    if link:
        st.link_button(f"Unlock Full Assessment — {price} →", link, use_container_width=True)
    st.markdown('<div class="divider-or">or</div>', unsafe_allow_html=True)
    st.markdown('''<div class="coupon-section">
      <div class="label">🎓 Writing Wives Skool Member? Scroll up and enter your coupon code above the form.</div>
    </div>''', unsafe_allow_html=True)

# ── URL verification ──────────────────────────────────────────────────────────
if not is_authenticated():
    sid = st.query_params.get("cover_session_id")
    if sid and not st.session_state.get("cover_stripe_checked"):
        st.session_state["cover_stripe_checked"] = True
        with st.spinner("Verifying your payment..."):
            ok, email = verify_stripe_session(sid)
        if ok:
            grant_access(reason="stripe", email=email)
            st.rerun()
        else:
            st.warning("Payment couldn't be verified yet — try refreshing in a moment.")

# ── Header ────────────────────────────────────────────────────────────────────
if st.button("← Back to All Tools", key="back_home"):
    st.switch_page("home.py")

_logo_b64 = get_logo_b64()
_logo_html = f'<div class="logo-wrap"><img src="data:image/png;base64,{_logo_b64}" alt="The Writing Wives"></div>' if _logo_b64 else ""
st.markdown(f"""
<div class="header-block">
  {_logo_html}
  <div class="eyebrow">Genre &amp; Trope AI Classroom</div>
  <h1>Cover Assessment</h1>
  <p class="sub">Upload your cover. Get a scored, genre-aware analysis — with specific fixes and a designer recommendation.</p>
</div>
""", unsafe_allow_html=True)

# ── Access banners ────────────────────────────────────────────────────────────
if is_authenticated():
    reason = st.session_state.get("cover_access_reason", "")
    email  = st.session_state.get("cover_access_email", "")
    if reason == "stripe":
        st.markdown(f'<div class="success-banner">✅ <strong>Payment confirmed{f" for {email}" if email else ""}.</strong> Full access unlocked.</div>', unsafe_allow_html=True)
        st.markdown('<div class="bookmark-tip">🔖 <strong>Bookmark this page</strong> — your URL contains your access token.</div>', unsafe_allow_html=True)
    elif reason in ("coupon", "lifetime"):
        st.markdown('<div class="success-banner">🎓 <strong>Access granted.</strong> Upload as many covers as you need.</div>', unsafe_allow_html=True)

# ── Coupon entry (shown when not yet authenticated) ──────────────────────────
if not is_authenticated():
    st.markdown('''
    <div class="coupon-section">
      <div class="label">🎓 Writing Wives Skool Member? Enter your coupon for free access.</div>
    </div>
    ''', unsafe_allow_html=True)
    if "cover_coupon_val" not in st.session_state:
        st.session_state["cover_coupon_val"] = ""
    coupon_col, btn_col = st.columns([3, 1])
    with coupon_col:
        st.text_input("Coupon code", placeholder="e.g. WRITINGWIVES",
                      label_visibility="collapsed", key="cover_coupon_val")
    with btn_col:
        apply_coupon = st.button("Apply →", key="cover_coupon_btn", use_container_width=True)
    if apply_coupon:
        entered = st.session_state.get("cover_coupon_val", "").strip()
        if entered and check_coupon(entered):
            grant_access(reason="coupon")
            st.rerun()
        else:
            st.error("That coupon code wasn't found. Double-check the spelling and try again.")

# ── Form ──────────────────────────────────────────────────────────────────────
with st.form("cover_form"):
    book_title = st.text_input("Book Title", placeholder="")
    col1, col2 = st.columns(2)
    with col1:
        genre = st.text_input("Genre / Sub-Genre", placeholder="")
    with col2:
        series = st.text_input("Series Name (optional)", placeholder="")
    cover_file = st.file_uploader(
        "Upload Your Cover",
        type=["jpg", "jpeg", "png", "webp"],
        help="Upload your book cover image (JPG, PNG, or WebP)"
    )
    submit_label = "Assess My Cover →" if is_authenticated() else "Get My Free Score →"
    submitted = st.form_submit_button(submit_label)

if submitted:
    if not cover_file:
        st.warning("Please upload your cover image before submitting.")
    else:
        # Show the cover they uploaded
        st.image(cover_file, caption=f"{book_title or 'Your Cover'}", width=280)
        image_bytes = cover_file.read()
        image_b64   = base64.b64encode(image_bytes).decode()
        image_mime  = cover_file.type or "image/jpeg"

        with st.spinner("Analysing your cover..."):
            try:
                result = call_cover_assessment(image_b64, genre, book_title, series, mime_type=image_mime)
                st.session_state["last_cover"] = {"result": result, "book_title": book_title}
            except json.JSONDecodeError:
                st.error("The AI returned an unexpected format. Please try again.")
                st.stop()
            except Exception as e:
                st.error("Something went wrong. Check your API key and try again.")
                st.exception(e)
                st.stop()

        st.success("Assessment complete!")
        st.divider()

        if is_authenticated():
            render_full_assessment(result, book_title)
        else:
            render_teaser(result)
            show_upgrade_card()

# ── If coupon just unlocked and we have cached result ─────────────────────────
elif is_authenticated() and st.session_state.get("last_cover") and st.session_state.get("cover_access_reason") == "coupon":
    prior = st.session_state["last_cover"]
    st.success("✅ Access granted! Here's your full assessment:")
    st.divider()
    render_full_assessment(prior["result"], prior["book_title"])
    del st.session_state["last_cover"]

# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="footer-note">
  The Writing Wives · Genre &amp; Trope AI Classroom ·
  <a href="https://thewritingwives.com" style="color:#999;">thewritingwives.com</a>
</div>
""", unsafe_allow_html=True)
